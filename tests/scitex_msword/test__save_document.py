#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-04 00:00:00
# File: tests/scitex_msword/test__save_document.py

"""Tests for ``scitex_msword._save_document.save_document``.

The point of the document-based save path (vs the writer-dict
``save_docx``) is that python-docx-level edits are preserved through
save. These tests verify the file lands at ``path``, that profile
``body_font`` / ``bold_font`` / ``body_font_size_pt`` / ``line_spacing``
hints are written into ``<w:docDefaults>``, that caller per-run
overrides survive, and that PRE/POST hooks fire in order with the
right context.

Style: AAA-marker comments on every test (STX-TQ002), one assertion per
test (STX-TQ007), ≥3-word descriptive names (STX-TQ003), no
``monkeypatch`` (PA-306 §3). State mutated by the sxm.hooks dispatcher
is reset by the autouse fixture below.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterator

import pytest


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@pytest.fixture(autouse=True)
def _clean_hooks_state() -> Iterator[None]:
    """Reset the hooks dispatcher around every test (no monkeypatch)."""
    from scitex_msword.hooks import _builtins, _dispatch, _lookup

    builtins_snapshot = dict(_builtins.ALL_HOOKS)
    iter_eps_original = _lookup._iter_entry_points
    _dispatch._reset()
    _lookup.reset()
    try:
        yield
    finally:
        _builtins.ALL_HOOKS.clear()
        _builtins.ALL_HOOKS.update(builtins_snapshot)
        _lookup._iter_entry_points = iter_eps_original
        _dispatch._reset()
        _lookup.reset()


def _find_docDefaults(doc):
    """Return the ``<w:docDefaults>`` element of ``doc.styles``."""
    from docx.oxml.ns import qn

    return doc.styles.element.find(qn("w:docDefaults"))


def _make_doc():
    pytest.importorskip("docx")
    from docx import Document

    return Document()


class TestSaveDocumentFileWrite:
    """``save_document`` writes a .docx file at the requested path."""

    def test_save_document_returns_pathlib_path_at_target(self, tmp_path):
        """The return value is the resolved ``Path`` of the output file."""
        # Arrange
        from scitex_msword import save_document

        doc = _make_doc()
        target = tmp_path / "out.docx"
        # Act
        returned = save_document(doc, target)
        # Assert
        assert returned == Path(target)

    def test_save_document_creates_nonempty_file_at_target(self, tmp_path):
        """The written ``.docx`` exists and is non-empty."""
        # Arrange
        from scitex_msword import save_document

        doc = _make_doc()
        target = tmp_path / "out.docx"
        # Act
        save_document(doc, target)
        # Assert
        assert target.exists() and target.stat().st_size > 0


class TestSaveDocumentProfileApplication:
    """When a profile is given, docDefaults are populated from its hints."""

    def test_boost_2026_profile_writes_mincho_to_docdefaults_ascii_slot(
        self, tmp_path
    ):
        """boost-2026 lands ＭＳ 明朝 in ``rPrDefault/rPr/rFonts @w:ascii``."""
        # Arrange
        from docx.oxml.ns import qn
        from scitex_msword import save_document

        doc = _make_doc()
        # Act
        save_document(doc, tmp_path / "out.docx", profile="boost-2026")
        # Assert
        r_fonts = _find_docDefaults(doc).find(
            qn("w:rPrDefault") + "/" + qn("w:rPr") + "/" + qn("w:rFonts")
        )
        assert r_fonts.get(qn("w:ascii")) == "ＭＳ 明朝"

    def test_boost_2026_profile_writes_mincho_to_docdefaults_hAnsi_slot(
        self, tmp_path
    ):
        """boost-2026 lands ＭＳ 明朝 in ``rPrDefault/rPr/rFonts @w:hAnsi``."""
        # Arrange
        from docx.oxml.ns import qn
        from scitex_msword import save_document

        doc = _make_doc()
        # Act
        save_document(doc, tmp_path / "out.docx", profile="boost-2026")
        # Assert
        r_fonts = _find_docDefaults(doc).find(
            qn("w:rPrDefault") + "/" + qn("w:rPr") + "/" + qn("w:rFonts")
        )
        assert r_fonts.get(qn("w:hAnsi")) == "ＭＳ 明朝"

    def test_boost_2026_profile_writes_mincho_to_docdefaults_eastAsia_slot(
        self, tmp_path
    ):
        """boost-2026 lands body_font (not bold_font) in ``docDefaults @w:eastAsia``.

        Per the v0.3.1 fix: docDefaults eastAsia carries the body
        typeface (Mincho), so non-bold Japanese body renders in
        Mincho. bold_font (Gothic) is applied at run level instead.
        """
        # Arrange
        from docx.oxml.ns import qn
        from scitex_msword import save_document

        doc = _make_doc()
        # Act
        save_document(doc, tmp_path / "out.docx", profile="boost-2026")
        # Assert
        r_fonts = _find_docDefaults(doc).find(
            qn("w:rPrDefault") + "/" + qn("w:rPr") + "/" + qn("w:rFonts")
        )
        assert r_fonts.get(qn("w:eastAsia")) == "ＭＳ 明朝"

    def test_boost_2026_profile_writes_10_5pt_body_font_size_in_half_points(
        self, tmp_path
    ):
        """boost-2026 lands 21 (= 10.5pt * 2 half-points) in ``rPrDefault/sz``."""
        # Arrange
        from docx.oxml.ns import qn
        from scitex_msword import save_document

        doc = _make_doc()
        # Act
        save_document(doc, tmp_path / "out.docx", profile="boost-2026")
        # Assert
        sz = _find_docDefaults(doc).find(
            qn("w:rPrDefault") + "/" + qn("w:rPr") + "/" + qn("w:sz")
        )
        assert sz.get(qn("w:val")) == "21"

    def test_boost_2026_profile_writes_240_line_spacing_for_single(
        self, tmp_path
    ):
        """boost-2026 lands ``spacing @w:line=240 @w:lineRule=auto`` (1.0 line)."""
        # Arrange
        from docx.oxml.ns import qn
        from scitex_msword import save_document

        doc = _make_doc()
        # Act
        save_document(doc, tmp_path / "out.docx", profile="boost-2026")
        # Assert
        spacing = _find_docDefaults(doc).find(
            qn("w:pPrDefault") + "/" + qn("w:pPr") + "/" + qn("w:spacing")
        )
        assert spacing.get(qn("w:line")) == "240"

    def test_none_profile_does_not_create_docdefaults_block(self, tmp_path):
        """``profile=None`` leaves ``docDefaults`` content untouched (no-op)."""
        # Arrange
        from docx.oxml.ns import qn
        from scitex_msword import save_document

        doc = _make_doc()
        styles_el = doc.styles.element
        before = styles_el.find(qn("w:docDefaults"))
        before_xml = etree_tostring(before) if before is not None else None
        # Act
        save_document(doc, tmp_path / "out.docx", profile=None)
        # Assert
        after = styles_el.find(qn("w:docDefaults"))
        after_xml = etree_tostring(after) if after is not None else None
        assert after_xml == before_xml


def etree_tostring(el):
    """Tiny helper for byte-level docDefaults equality."""
    from lxml import etree

    return etree.tostring(el)


class TestSaveDocumentBoldFontRunOverride:
    """bold_font is applied at run level to bold runs only (v0.3.1 fix)."""

    def test_bold_run_gets_bold_font_in_eastAsia_slot(self, tmp_path):
        """A run with ``run.bold=True`` ends up with eastAsia=bold_font."""
        # Arrange
        from docx.oxml.ns import qn
        from scitex_msword import save_document

        doc = _make_doc()
        para = doc.add_paragraph()
        run = para.add_run("ボールド文字列")
        run.bold = True
        # Act
        save_document(doc, tmp_path / "out.docx", profile="boost-2026")
        # Assert
        rpr = run._r.find(qn("w:rPr"))
        rfonts = rpr.find(qn("w:rFonts"))
        assert rfonts.get(qn("w:eastAsia")) == "ＭＳ ゴシック"

    def test_non_bold_run_does_not_get_bold_font_in_eastAsia(self, tmp_path):
        """A run with ``run.bold=False`` keeps eastAsia at docDefaults (Mincho)."""
        # Arrange
        from docx.oxml.ns import qn
        from scitex_msword import save_document

        doc = _make_doc()
        para = doc.add_paragraph()
        run = para.add_run("通常本文")
        run.bold = False
        # Act
        save_document(doc, tmp_path / "out.docx", profile="boost-2026")
        # Assert
        rpr = run._r.find(qn("w:rPr"))
        rfonts = None if rpr is None else rpr.find(qn("w:rFonts"))
        # Either no rFonts element on the run (inheriting docDefaults Mincho)
        # OR rFonts exists but eastAsia is unset / not Gothic.
        eastAsia = None if rfonts is None else rfonts.get(qn("w:eastAsia"))
        assert eastAsia != "ＭＳ ゴシック"


class TestSaveDocumentHookChain:
    """``save_document`` runs PRE_SAVE / POST_SAVE in order."""

    def test_pre_save_hook_runs_before_file_is_written(self, tmp_path):
        """PRE_SAVE hooks see a doc that has not yet been written to disk."""
        # Arrange
        from scitex_msword import save_document
        from scitex_msword.hooks import Hook, Phase, register

        saw_before = []
        out = tmp_path / "out.docx"

        def pre(doc, ctx):
            saw_before.append(out.exists())

        register(
            Hook(
                id="SXM-T-PRE",
                phase=Phase.PRE_SAVE,
                severity="info",
                category="t",
                message="m",
                suggestion="s",
                fn=pre,
            )
        )
        doc = _make_doc()
        # Act
        save_document(doc, out, profile=None)
        # Assert
        assert saw_before == [False]

    def test_post_save_hook_receives_out_path_after_file_is_written(
        self, tmp_path
    ):
        """POST_SAVE hooks see a written file at ``out_path``."""
        # Arrange
        from scitex_msword import save_document
        from scitex_msword.hooks import Hook, Phase, register

        saw_after = []
        out = tmp_path / "out.docx"

        def post(doc, ctx, *, out_path):
            saw_after.append((out_path, out_path.exists()))

        register(
            Hook(
                id="SXM-T-POST",
                phase=Phase.POST_SAVE,
                severity="info",
                category="t",
                message="m",
                suggestion="s",
                fn=post,
            )
        )
        doc = _make_doc()
        # Act
        save_document(doc, out, profile=None)
        # Assert
        assert saw_after == [(out, True)]


# EOF
