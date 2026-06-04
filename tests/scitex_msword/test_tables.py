#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-04 12:00:00
# File: tests/scitex_msword/test_tables.py

"""Tests for scitex_msword.tables.insert_table_after_paragraph."""

import io

import pytest

pytestmark = pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="python-docx not installed",
)


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


def _doc_with_paragraphs(texts):
    """Build an in-memory python-docx Document with one paragraph per text."""
    import docx

    doc = docx.Document()
    for text in texts:
        doc.add_paragraph(text)
    return doc


def _roundtrip(doc):
    """Save to BytesIO then re-load, simulating disk round-trip."""
    import docx

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return docx.Document(buf)


# ---------------------------------------------------------------------------
# Import surface
# ---------------------------------------------------------------------------


class TestImportSurface:
    """Smoke tests for module imports."""

    def test_module_is_importable(self):
        """scitex_msword.tables should import cleanly."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.tables")
        # Assert
        assert callable(mod.insert_table_after_paragraph)

    def test_top_level_export_present(self):
        """scitex_msword top-level should re-export insert_table_after_paragraph."""
        # Arrange
        import scitex_msword as sxm

        # Act
        attr = getattr(sxm, "insert_table_after_paragraph", None)
        # Assert
        assert callable(attr)


# ---------------------------------------------------------------------------
# Element structure — pure build, no Track Changes
# ---------------------------------------------------------------------------


class TestTableElementShape:
    """The freshly inserted <w:tbl> carries the expected OOXML scaffolding."""

    def test_returns_tbl_element(self):
        """Return value should be the inserted <w:tbl> element."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc, 0, [("h1", "h2"), ("a", "b")]
        )

        # Assert
        assert tbl.tag == _w("tbl")

    def test_inserts_after_anchor_paragraph(self):
        """The new <w:tbl> should sit as the next sibling of paragraphs[idx]."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["before", "anchor", "after"])

        # Act
        insert_table_after_paragraph(doc, 1, [("h1", "h2"), ("a", "b")])
        anchor_el = doc.paragraphs[1]._p
        sibling = anchor_el.getnext()

        # Assert
        assert sibling.tag == _w("tbl")

    def test_emits_one_row_per_input_row(self):
        """Generated <w:tr> count should match len(rows)."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("h1", "h2"), ("r1c1", "r1c2"), ("r2c1", "r2c2")],
        )
        row_count = len(tbl.findall(_w("tr")))

        # Assert
        assert row_count == 3

    def test_emits_one_gridcol_per_column(self):
        """<w:tblGrid> should have one <w:gridCol> per width entry."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("a", "b", "c")],
            col_widths_dxa=(2000, 3000, 4000),
            header_row=False,
        )
        grid_cols = tbl.find(_w("tblGrid")).findall(_w("gridCol"))

        # Assert
        assert len(grid_cols) == 3

    def test_records_column_widths_in_dxa(self):
        """Each <w:gridCol w:w="…"> should equal the input col_widths_dxa entry."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("a", "b")],
            col_widths_dxa=(3000, 6000),
            header_row=False,
        )
        widths = [gc.get(_w("w")) for gc in tbl.find(_w("tblGrid"))]

        # Assert
        assert widths == ["3000", "6000"]

    def test_cell_text_appears_in_first_row(self):
        """The cell text of row 0 / column 0 should round-trip into <w:t>."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc, 0, [("役割", "モジュール名")], header_row=True
        )
        first_t = tbl.find(_w("tr")).find(_w("tc")).find(_w("p")).find(_w("r")).find(_w("t"))

        # Assert
        assert first_t.text == "役割"

    def test_header_row_uses_header_font(self):
        """Row 0 should be styled with the header_font when header_row=True."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("役割", "モジュール名"), ("a", "b")],
            header_font="MS ゴシック",
            body_font="MS 明朝",
            header_row=True,
        )
        header_tc = tbl.find(_w("tr")).find(_w("tc"))
        rFonts = (
            header_tc.find(_w("p")).find(_w("r")).find(_w("rPr")).find(_w("rFonts"))
        )

        # Assert
        assert rFonts.get(_w("eastAsia")) == "MS ゴシック"

    def test_body_row_uses_body_font(self):
        """Row 1 should be styled with the body_font when header_row=True."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("h1", "h2"), ("a", "b")],
            header_font="MS ゴシック",
            body_font="MS 明朝",
            header_row=True,
        )
        body_tc = tbl.findall(_w("tr"))[1].find(_w("tc"))
        rFonts = (
            body_tc.find(_w("p")).find(_w("r")).find(_w("rPr")).find(_w("rFonts"))
        )

        # Assert
        assert rFonts.get(_w("eastAsia")) == "MS 明朝"

    def test_header_row_runs_are_bold(self):
        """Header row runs should carry <w:b/>."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc, 0, [("h1", "h2"), ("a", "b")], header_row=True
        )
        header_rPr = (
            tbl.find(_w("tr")).find(_w("tc")).find(_w("p")).find(_w("r")).find(_w("rPr"))
        )

        # Assert
        assert header_rPr.find(_w("b")) is not None

    def test_body_row_runs_are_not_bold(self):
        """Body row runs should not carry <w:b/>."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc, 0, [("h1", "h2"), ("a", "b")], header_row=True
        )
        body_rPr = (
            tbl.findall(_w("tr"))[1]
            .find(_w("tc"))
            .find(_w("p"))
            .find(_w("r"))
            .find(_w("rPr"))
        )

        # Assert
        assert body_rPr.find(_w("b")) is None

    def test_font_size_pt_to_half_points(self):
        """font_size_pt=10.5 should write <w:sz w:val="21">."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("a", "b")],
            header_row=False,
            font_size_pt=10.5,
        )
        sz = (
            tbl.find(_w("tr"))
            .find(_w("tc"))
            .find(_w("p"))
            .find(_w("r"))
            .find(_w("rPr"))
            .find(_w("sz"))
        )

        # Assert
        assert sz.get(_w("val")) == "21"


# ---------------------------------------------------------------------------
# Round-trip — document save + reopen still sees the table
# ---------------------------------------------------------------------------


class TestRoundtripPreservesTable:
    """After save/load the document should still carry the inserted table."""

    def test_doc_tables_count_increases_by_one(self):
        """doc.tables length should grow by exactly one after insertion."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])
        before = len(doc.tables)

        # Act
        insert_table_after_paragraph(
            doc, 0, [("h1", "h2"), ("a", "b")]
        )
        reloaded = _roundtrip(doc)

        # Assert
        assert len(reloaded.tables) == before + 1

    def test_roundtrip_preserves_cell_text(self):
        """Reading back via python-docx should see the same cell text."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        insert_table_after_paragraph(
            doc, 0, [("役割", "モジュール"), ("論文執筆", "scitex-writer")]
        )
        reloaded = _roundtrip(doc)
        first_row_texts = [c.text for c in reloaded.tables[-1].rows[0].cells]

        # Assert
        assert first_row_texts == ["役割", "モジュール"]


# ---------------------------------------------------------------------------
# Validation — boundary conditions raise clearly
# ---------------------------------------------------------------------------


class TestInputValidation:
    """Bad inputs should raise IndexError/ValueError, not silently misbehave."""

    def test_paragraph_index_out_of_range_raises(self):
        """paragraph_index past end-of-doc should raise IndexError."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["only"])
        ctx = pytest.raises(IndexError)

        # Act
        # Assert
        with ctx:
            insert_table_after_paragraph(doc, 99, [("a", "b")])

    def test_empty_rows_raises_value_error(self):
        """rows=[] should raise ValueError."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])
        ctx = pytest.raises(ValueError)

        # Act
        # Assert
        with ctx:
            insert_table_after_paragraph(doc, 0, [])

    def test_uneven_columns_raises_value_error(self):
        """Inconsistent column counts across rows should raise ValueError."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])
        ctx = pytest.raises(ValueError)

        # Act
        # Assert
        with ctx:
            insert_table_after_paragraph(doc, 0, [("a", "b"), ("c",)])

    def test_col_widths_mismatch_raises_value_error(self):
        """col_widths_dxa length mismatch with row width should raise ValueError."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])
        ctx = pytest.raises(ValueError)

        # Act
        # Assert
        with ctx:
            insert_table_after_paragraph(
                doc, 0, [("a", "b")], col_widths_dxa=(3000,)
            )


# ---------------------------------------------------------------------------
# Track-changes integration
# ---------------------------------------------------------------------------


class TestTrackChangesIntegration:
    """Row-level <w:trPr><w:ins/></w:trPr> emission rules."""

    def test_no_wrap_when_tc_off_and_autoexplicit(self):
        """Default (track_changes=None) on a TC-off doc emits NO trPr/ins markers."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])  # default: TC off

        # Act
        tbl = insert_table_after_paragraph(
            doc, 0, [("h1", "h2"), ("a", "b")]
        )
        ins_count = len(tbl.findall(f"{_w('tr')}/{_w('trPr')}/{_w('ins')}"))

        # Assert
        assert ins_count == 0

    def test_force_wrap_marks_every_row(self):
        """track_changes=True should mark every row with <w:trPr><w:ins/>."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("h1", "h2"), ("a", "b"), ("c", "d")],
            track_changes=True,
            track_changes_author="claude-agent",
        )
        ins_count = len(tbl.findall(f"{_w('tr')}/{_w('trPr')}/{_w('ins')}"))

        # Assert
        assert ins_count == 3

    def test_force_no_wrap_overrides_tc_on_doc(self):
        """track_changes=False should suppress markers even on a TC-on doc."""
        # Arrange
        from scitex_msword import (
            enable_track_changes,
            insert_table_after_paragraph,
        )

        doc = _doc_with_paragraphs(["anchor"])
        enable_track_changes(doc, enabled=True)

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("h1", "h2"), ("a", "b")],
            track_changes=False,
        )
        ins_count = len(tbl.findall(f"{_w('tr')}/{_w('trPr')}/{_w('ins')}"))

        # Assert
        assert ins_count == 0

    def test_auto_detect_wraps_when_tc_on(self):
        """track_changes=None on a TC-on doc should auto-wrap each row."""
        # Arrange
        from scitex_msword import (
            enable_track_changes,
            insert_table_after_paragraph,
        )

        doc = _doc_with_paragraphs(["anchor"])
        enable_track_changes(doc, enabled=True)

        # Act
        tbl = insert_table_after_paragraph(
            doc, 0, [("h1", "h2"), ("a", "b")]
        )
        ins_count = len(tbl.findall(f"{_w('tr')}/{_w('trPr')}/{_w('ins')}"))

        # Assert
        assert ins_count == 2

    def test_ins_author_attr_matches_argument(self):
        """track_changes_author should land on every w:ins/@w:author."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("a", "b"), ("c", "d")],
            track_changes=True,
            track_changes_author="claude-agent",
        )
        ins_elements = tbl.findall(f"{_w('tr')}/{_w('trPr')}/{_w('ins')}")
        authors = {el.get(_w("author")) for el in ins_elements}

        # Assert
        assert authors == {"claude-agent"}

    def test_ins_date_attr_matches_argument(self):
        """track_changes_date should land verbatim on every w:ins/@w:date."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])
        iso = "2026-06-04T12:00:00Z"

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("a", "b"), ("c", "d")],
            track_changes=True,
            track_changes_date=iso,
        )
        ins_elements = tbl.findall(f"{_w('tr')}/{_w('trPr')}/{_w('ins')}")
        dates = {el.get(_w("date")) for el in ins_elements}

        # Assert
        assert dates == {iso}

    def test_ins_ids_are_contiguous_integers(self):
        """w:id values across rows should be a contiguous, increasing run."""
        # Arrange
        from scitex_msword import insert_table_after_paragraph

        doc = _doc_with_paragraphs(["anchor"])

        # Act
        tbl = insert_table_after_paragraph(
            doc,
            0,
            [("a", "b"), ("c", "d"), ("e", "f")],
            track_changes=True,
        )
        ins_elements = tbl.findall(f"{_w('tr')}/{_w('trPr')}/{_w('ins')}")
        ids = [int(el.get(_w("id"))) for el in ins_elements]

        # Assert
        assert ids == [ids[0], ids[0] + 1, ids[0] + 2]


# ---------------------------------------------------------------------------
# CLI smoke test
# ---------------------------------------------------------------------------


class TestCliInsertTable:
    """The scitex-msword insert-table CLI should produce a usable .docx."""

    def test_cli_creates_output_file(self, tmp_path):
        """Running cli.main on insert-table should leave an output file with a table."""
        # Arrange
        import docx as _docx

        from scitex_msword.cli import main

        src = tmp_path / "src.docx"
        out = tmp_path / "out.docx"
        seed_doc = _docx.Document()
        seed_doc.add_paragraph("anchor")
        seed_doc.save(str(src))

        # Act
        rc = main(
            [
                "insert-table",
                "--path",
                str(src),
                "--out",
                str(out),
                "--paragraph-index",
                "0",
                "--rows",
                '[["役割","モジュール名"],["論文執筆","scitex-writer"]]',
            ]
        )
        reloaded = _docx.Document(str(out))

        # Assert
        assert rc == 0 and len(reloaded.tables) == 1
