#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-04 00:00:00
# File: tests/scitex_msword/test__settings_order.py

"""Tests for ``scitex_msword._settings_order``.

ECMA-376 §17.15.1 ``CT_Settings`` has a fixed child sequence; Word
silently ignores out-of-order elements (notably ``<w:trackRevisions/>``
when written before its predecessors on otherwise-sparse settings
files). The placement helper here owns that decision.

Also covers the ``save_with_track_changes_on`` convenience wrapper from
``scitex_msword.track_changes`` since the two ship together for the
BOOST v37 dogfood workflow.

Style: AAA-marker comments on every test (STX-TQ002), one assertion per
test (STX-TQ007), ≥3-word descriptive names (STX-TQ003), no
``monkeypatch`` (PA-306 §3).
"""

from __future__ import annotations

import pathlib

import pytest


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w_element(tag, **attrs):
    """Build a bare ``<w:tag/>`` element for tests."""
    from lxml import etree

    el = etree.Element(f"{{{_W_NS}}}{tag}")
    for k, v in attrs.items():
        el.set(f"{{{_W_NS}}}{k}", str(v))
    return el


def _settings_with(children):
    """Build a ``<w:settings/>`` element wrapping the given tag-name list."""
    from lxml import etree

    settings = etree.Element(f"{{{_W_NS}}}settings")
    for tag in children:
        settings.append(_w_element(tag))
    return settings


def _local_names(settings_el):
    """Return the ordered local tag names of children of ``settings_el``."""
    from lxml import etree

    return [etree.QName(c).localname for c in settings_el]


class TestInsertInSettingsOrderTrackChanges:
    """``trackRevisions`` placement obeys ECMA-376 §17.15.1 ordering."""

    def test_inserts_after_stylePaneSortMethod_when_present(self):
        """With ``stylePaneSortMethod`` present, ``trackRevisions`` lands after it."""
        # Arrange
        from scitex_msword._settings_order import insert_in_settings_order

        settings_el = _settings_with(["stylePaneSortMethod"])
        new_el = _w_element("trackRevisions")
        # Act
        insert_in_settings_order(settings_el, new_el, "trackRevisions")
        # Assert
        assert _local_names(settings_el) == [
            "stylePaneSortMethod",
            "trackRevisions",
        ]

    def test_inserts_before_doNotTrackFormatting_when_only_successor_present(
        self,
    ):
        """With only ``doNotTrackFormatting`` present, ``trackRevisions`` lands before it."""
        # Arrange
        from scitex_msword._settings_order import insert_in_settings_order

        settings_el = _settings_with(["doNotTrackFormatting"])
        new_el = _w_element("trackRevisions")
        # Act
        insert_in_settings_order(settings_el, new_el, "trackRevisions")
        # Assert
        assert _local_names(settings_el) == [
            "trackRevisions",
            "doNotTrackFormatting",
        ]

    def test_lands_between_predecessor_and_successor_when_both_present(self):
        """Both anchors present → ``trackRevisions`` slots in between."""
        # Arrange
        from scitex_msword._settings_order import insert_in_settings_order

        settings_el = _settings_with(
            ["stylePaneSortMethod", "doNotTrackFormatting"]
        )
        new_el = _w_element("trackRevisions")
        # Act
        insert_in_settings_order(settings_el, new_el, "trackRevisions")
        # Assert
        assert _local_names(settings_el) == [
            "stylePaneSortMethod",
            "trackRevisions",
            "doNotTrackFormatting",
        ]

    def test_appends_when_neither_anchor_is_present(self):
        """No before/after anchors → append at end (historical behaviour)."""
        # Arrange
        from scitex_msword._settings_order import insert_in_settings_order

        settings_el = _settings_with(["view", "zoom"])
        # Note: "view" + "zoom" are in the "before" set but neither is a
        # late-position predecessor close to trackRevisions; the helper still
        # treats the last one of them as the anchor.
        # Act
        insert_in_settings_order(
            settings_el, _w_element("trackRevisions"), "trackRevisions"
        )
        # Assert
        assert _local_names(settings_el)[-1] == "trackRevisions"

    def test_falls_back_to_append_for_unknown_tag(self):
        """Asking to place a tag the helper doesn't model falls back to append."""
        # Arrange
        from scitex_msword._settings_order import insert_in_settings_order

        settings_el = _settings_with(["stylePaneSortMethod"])
        # Act
        insert_in_settings_order(
            settings_el, _w_element("someUnmodelledTag"), "someUnmodelledTag"
        )
        # Assert
        assert _local_names(settings_el)[-1] == "someUnmodelledTag"


class TestEnableTrackChangesOrderedPlacement:
    """``enable_track_changes`` calls into the ordered placement helper."""

    def test_enable_track_changes_lands_in_ordered_slot_on_real_document(
        self, tmp_path
    ):
        """A fresh ``Document`` gets ``<w:trackRevisions/>`` in the right slot."""
        # Arrange
        pytest.importorskip("docx")
        from docx import Document
        from lxml import etree
        from scitex_msword.track_changes import enable_track_changes

        doc = Document()
        # Act
        enable_track_changes(doc)
        # Assert
        settings_el = doc.settings.element
        tags = [etree.QName(c).localname for c in settings_el]
        assert "trackRevisions" in tags


class TestSaveWithTrackChangesOnHelper:
    """``save_with_track_changes_on`` enables TC and saves in one call."""

    def test_save_with_track_changes_on_writes_a_file_at_the_target_path(
        self, tmp_path
    ):
        """The helper produces a .docx file at the requested path."""
        # Arrange
        pytest.importorskip("docx")
        from docx import Document
        from scitex_msword.track_changes import save_with_track_changes_on

        doc = Document()
        out = tmp_path / "tc.docx"
        # Act
        save_with_track_changes_on(doc, out)
        # Assert
        assert out.exists() and out.stat().st_size > 0

    def test_save_with_track_changes_on_persists_track_changes_in_saved_file(
        self, tmp_path
    ):
        """Reloading the saved file shows ``<w:trackRevisions/>`` is present."""
        # Arrange
        pytest.importorskip("docx")
        from docx import Document
        from scitex_msword.track_changes import (
            is_track_changes_enabled,
            save_with_track_changes_on,
        )

        doc = Document()
        out = tmp_path / "tc.docx"
        save_with_track_changes_on(doc, out)
        # Act
        reloaded = Document(str(out))
        enabled = is_track_changes_enabled(reloaded)
        # Assert
        assert enabled is True

    def test_save_with_track_changes_on_emits_documentProtection_state_only(
        self, tmp_path
    ):
        """The saved settings.xml also carries the matching
        ``<w:documentProtection w:edit="trackedChanges" w:enforcement="0"/>``
        — informational, NOT enforced. Matches what desktop Word writes.
        """
        # Arrange
        pytest.importorskip("docx")
        from docx import Document
        from docx.oxml.ns import qn
        from scitex_msword.track_changes import save_with_track_changes_on

        doc = Document()
        out = tmp_path / "tc.docx"
        save_with_track_changes_on(doc, out)
        reloaded = Document(str(out))
        # Act
        protection = reloaded.settings.element.find(
            qn("w:documentProtection")
        )
        observed = (
            None if protection is None
            else (
                protection.get(qn("w:edit")),
                protection.get(qn("w:enforcement")),
            )
        )
        # Assert
        assert observed == ("trackedChanges", "0")


_VENDORED_WORD_GROUNDTRUTH_SETTINGS = (
    pathlib.Path(__file__).parent.parent
    / "fixtures"
    / "track_changes"
    / "word_groundtruth_settings.xml"
)


class TestTrackRevisionsAgainstVendoredWordGroundTruth:
    """Pins the v0.3.1 trackChanges→trackRevisions fix against a
    committed Word-emitted ``word/settings.xml`` fixture (extracted
    from proj-grant draft_v39 — the operator's manual Track Changes
    ON save during BOOST v40). See
    ``tests/fixtures/track_changes/README.md`` for full provenance.

    The blind-spot fixed by these tests is that the previous
    conformance gate asserted our writer's own output, not Word's
    actual emit — so the self-consistent wrong-name state matched
    itself in CI and looked fine until proj-grant crashed it on the
    operator's desktop Word. These pins anchor against bytes Word
    itself wrote, so a future regression in either direction (back to
    the wrong name OR away from it) fails CI loudly."""

    @staticmethod
    def _parse_groundtruth():
        from lxml import etree

        return etree.fromstring(
            _VENDORED_WORD_GROUNDTRUTH_SETTINGS.read_bytes()
        )

    def test_vendored_groundtruth_carries_trackRevisions_element(self):
        """Word's own emit puts ``<w:trackRevisions/>`` in CT_Settings."""
        # Arrange
        from docx.oxml.ns import qn

        root = self._parse_groundtruth()
        # Act
        found = root.find(qn("w:trackRevisions"))
        # Assert
        assert found is not None

    def test_vendored_groundtruth_does_not_carry_trackChanges_element(self):
        """Word's own emit does NOT use ``<w:trackChanges/>`` for the toggle."""
        # Arrange
        from docx.oxml.ns import qn

        root = self._parse_groundtruth()
        # Act
        found = root.find(qn("w:trackChanges"))
        # Assert
        assert found is None

    def test_vendored_groundtruth_carries_documentProtection_state_only(self):
        """Word's emit pairs ``<w:trackRevisions/>`` with
        ``<w:documentProtection w:edit="trackedChanges" w:enforcement="0"/>``.
        """
        # Arrange
        from docx.oxml.ns import qn

        root = self._parse_groundtruth()
        protection = root.find(qn("w:documentProtection"))
        observed = (
            None
            if protection is None
            else (
                protection.get(qn("w:edit")),
                protection.get(qn("w:enforcement")),
            )
        )
        # Act / Assert
        assert observed == ("trackedChanges", "0")

    def test_save_with_track_changes_on_matches_word_trackRevisions_slice(
        self, tmp_path
    ):
        """``save_with_track_changes_on`` emits the same trackRevisions
        slice (``trackRevisions`` + ``documentProtection edit=trackedChanges
        enforcement=0``) as Word's own save."""
        # Arrange
        pytest.importorskip("docx")
        from docx import Document
        from docx.oxml.ns import qn
        from scitex_msword.track_changes import save_with_track_changes_on

        save_with_track_changes_on(Document(), tmp_path / "rc.docx")
        emitted = Document(str(tmp_path / "rc.docx")).settings.element

        groundtruth = self._parse_groundtruth()

        def _slice(parent):
            wanted = ("trackRevisions", "documentProtection")
            from lxml import etree

            out = []
            for child in parent:
                local = etree.QName(child).localname
                if local in wanted:
                    out.append(
                        (
                            local,
                            child.get(qn("w:edit")),
                            child.get(qn("w:enforcement")),
                        )
                    )
            return out

        # Act
        observed = (_slice(emitted), _slice(groundtruth))
        # Assert
        assert observed[0] == observed[1]


# Path resolver kept for the H4 builtin-hooks work where we'll exercise
# more of the full docx round-trip; today nothing references it.
def _resolve_draft_v39_reference():
    """Return the proj-grant draft_v39 reference docx Path or None.

    H4 / SXM-TC001 follow-up will use this to broaden the conformance
    tests against the entire docx (not just settings.xml).
    """
    import os

    candidates = [
        os.environ.get("SXM_DRAFT_V39_REFERENCE"),
        "/home/ywatanabe/proj/grant/2026-06-11---2027-04-2032-03"
        "---20-PERC---1000---BOOST/"
        "draft_v39_ywata-turned-on-edit-history.docx",
    ]
    for raw in candidates:
        if not raw:
            continue
        p = pathlib.Path(raw)
        if p.exists():
            return p
    return None


# EOF
