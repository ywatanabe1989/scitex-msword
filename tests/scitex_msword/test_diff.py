#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-02 00:00:00
# File: tests/scitex_msword/test_diff.py

"""Tests for scitex_msword.diff module."""

from pathlib import Path

import pytest

pytestmark = pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="python-docx not installed",
)


def _build_docx(tmp_path: Path, name: str, paragraphs):
    """Create a tiny docx with the given list of (text, bold) tuples."""
    import docx

    doc = docx.Document()
    for text, bold in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
    out = tmp_path / name
    doc.save(str(out))
    return out


class TestDiffDocxImport:
    """Smoke tests for diff_docx import surface."""

    def test_diff_docx_is_importable_from_module(self):
        """diff_docx should be importable from scitex_msword.diff."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.diff")
        # Assert
        assert callable(mod.diff_docx)

    def test_summarize_diff_is_importable_from_module(self):
        """summarize_diff should be importable from scitex_msword.diff."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.diff")
        # Assert
        assert callable(mod.summarize_diff)


class TestDiffDocxIdentical:
    """Two identical docs should produce only 'equal' ops."""

    def test_identical_docs_yield_only_equal_ops(self, tmp_path):
        """All ops should be 'equal' when comparing a file with itself."""
        # Arrange
        from scitex_msword.diff import diff_docx

        path = _build_docx(
            tmp_path,
            "same.docx",
            [("Hello world", False), ("Second line", False)],
        )
        # Act
        ops = diff_docx(path, path)
        # Assert
        assert all(op["op"] == "equal" for op in ops)

    def test_identical_docs_yield_one_op_per_paragraph(self, tmp_path):
        """The number of 'equal' ops should equal the paragraph count."""
        # Arrange
        from scitex_msword.diff import diff_docx

        path = _build_docx(
            tmp_path,
            "same2.docx",
            [("p1", False), ("p2", False), ("p3", False)],
        )
        # Act
        ops = diff_docx(path, path)
        # Assert
        assert len(ops) == 3


class TestDiffDocxInsert:
    """A paragraph added to the right doc should appear as an 'insert' op."""

    def test_insert_op_detected_when_b_has_extra_paragraph(self, tmp_path):
        """diff_docx should emit an 'insert' op for the new paragraph."""
        # Arrange
        from scitex_msword.diff import diff_docx

        a = _build_docx(tmp_path, "a.docx", [("first", False)])
        b = _build_docx(
            tmp_path,
            "b.docx",
            [("first", False), ("brand new", False)],
        )
        # Act
        ops = diff_docx(a, b)
        # Assert
        assert any(op["op"] == "insert" and op["text_b"] == "brand new" for op in ops)


class TestDiffDocxDelete:
    """A paragraph removed in the right doc should appear as a 'delete' op."""

    def test_delete_op_detected_when_b_drops_paragraph(self, tmp_path):
        """diff_docx should emit a 'delete' op for the dropped paragraph."""
        # Arrange
        from scitex_msword.diff import diff_docx

        a = _build_docx(
            tmp_path,
            "a.docx",
            [("first", False), ("to be dropped", False)],
        )
        b = _build_docx(tmp_path, "b.docx", [("first", False)])
        # Act
        ops = diff_docx(a, b)
        # Assert
        assert any(
            op["op"] == "delete" and op["text_a"] == "to be dropped" for op in ops
        )


class TestDiffDocxModify:
    """A paragraph with different text should appear as a 'modify' op."""

    def test_modify_op_detected_when_paragraph_text_changes(self, tmp_path):
        """A paragraph text edit should produce a 'modify' op."""
        # Arrange
        from scitex_msword.diff import diff_docx

        a = _build_docx(tmp_path, "a.docx", [("original text", False)])
        b = _build_docx(tmp_path, "b.docx", [("edited text", False)])
        # Act
        ops = diff_docx(a, b)
        # Assert
        assert any(op["op"] == "modify" for op in ops)

    def test_modify_op_exposes_text_a(self, tmp_path):
        """A 'modify' op should expose text_a for the changed paragraph."""
        # Arrange
        from scitex_msword.diff import diff_docx

        a = _build_docx(tmp_path, "a.docx", [("original text", False)])
        b = _build_docx(tmp_path, "b.docx", [("edited text", False)])
        # Act
        ops = diff_docx(a, b)
        modify_ops = [op for op in ops if op["op"] == "modify"]
        # Assert
        assert modify_ops[0]["text_a"] == "original text"

    def test_modify_op_exposes_text_b(self, tmp_path):
        """A 'modify' op should expose text_b for the changed paragraph."""
        # Arrange
        from scitex_msword.diff import diff_docx

        a = _build_docx(tmp_path, "a.docx", [("original text", False)])
        b = _build_docx(tmp_path, "b.docx", [("edited text", False)])
        # Act
        ops = diff_docx(a, b)
        modify_ops = [op for op in ops if op["op"] == "modify"]
        # Assert
        assert modify_ops[0]["text_b"] == "edited text"


class TestDiffDocxRunDelta:
    """Bold-only changes should be captured in runs_changed."""

    def test_bold_change_produces_modify_op(self, tmp_path):
        """A bold toggle on changed text should produce a modify op."""
        # Arrange
        from scitex_msword.diff import diff_docx

        a = _build_docx(tmp_path, "a.docx", [("same text", False)])
        b = _build_docx(tmp_path, "b.docx", [("same text changed", True)])
        # Act
        ops = diff_docx(a, b)
        modify_ops = [op for op in ops if op["op"] == "modify"]
        # Assert
        assert modify_ops

    def test_bold_change_modify_op_carries_runs_changed_key(self, tmp_path):
        """The modify op for a bold-toggled paragraph should expose runs_changed."""
        # Arrange
        from scitex_msword.diff import diff_docx

        a = _build_docx(tmp_path, "a.docx", [("same text", False)])
        b = _build_docx(tmp_path, "b.docx", [("same text changed", True)])
        # Act
        ops = diff_docx(a, b)
        modify_ops = [op for op in ops if op["op"] == "modify"]
        # Assert
        assert "runs_changed" in modify_ops[0]


class TestSummarizeDiff:
    """summarize_diff should count op types."""

    def test_summarize_diff_counts_equal_ops(self, tmp_path):
        """summarize_diff should report the 'equal' op count."""
        # Arrange
        from scitex_msword.diff import diff_docx, summarize_diff

        path = _build_docx(tmp_path, "same.docx", [("p1", False), ("p2", False)])
        ops = diff_docx(path, path)
        # Act
        summary = summarize_diff(ops)
        # Assert
        assert summary["equal"] == 2

    def test_summarize_diff_counts_insert_ops(self, tmp_path):
        """summarize_diff should report the 'insert' op count."""
        # Arrange
        from scitex_msword.diff import diff_docx, summarize_diff

        a = _build_docx(tmp_path, "a.docx", [("p1", False)])
        b = _build_docx(
            tmp_path, "b.docx", [("p1", False), ("new", False), ("new2", False)]
        )
        ops = diff_docx(a, b)
        # Act
        summary = summarize_diff(ops)
        # Assert
        assert summary["insert"] == 2


class TestDiffDocxErrors:
    """diff_docx should raise on missing/invalid inputs."""

    def test_diff_docx_raises_for_nonexistent_path(self, tmp_path):
        """diff_docx should raise when the input file does not exist."""
        # Arrange
        from scitex_msword.diff import diff_docx

        good = _build_docx(tmp_path, "ok.docx", [("hello", False)])
        bad = tmp_path / "does_not_exist.docx"
        ctx = pytest.raises(Exception)
        # Act
        # Assert
        with ctx:
            diff_docx(good, bad)


if __name__ == "__main__":
    import os

    pytest.main([os.path.abspath(__file__)])
