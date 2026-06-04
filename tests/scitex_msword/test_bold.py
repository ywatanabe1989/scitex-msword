#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-02 00:00:00
# File: tests/scitex_msword/test_bold.py

"""Tests for scitex_msword.bold module."""

import pytest

pytestmark = pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="python-docx not installed",
)


def _doc_with_text(paragraphs):
    """Build a Document with one run per paragraph holding the supplied text."""
    import docx

    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


class TestImportSurface:
    """Smoke tests."""

    def test_preserve_bold_tokens_is_importable(self):
        """preserve_bold_tokens should be importable from the module."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.bold")
        # Assert
        assert callable(mod.preserve_bold_tokens)


class TestPreserveBoldTokensSplitting:
    """preserve_bold_tokens should split runs around token hits."""

    def test_single_token_in_paragraph_creates_three_runs(self):
        """'foo BAR baz' with token 'BAR' should yield 3 runs."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["foo BAR baz"])
        # Act
        preserve_bold_tokens(doc, ["BAR"])
        # Assert
        assert len(doc.paragraphs[0].runs) == 3

    def test_token_run_is_bold(self):
        """The run corresponding to the token should be bold."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["foo BAR baz"])
        # Act
        preserve_bold_tokens(doc, ["BAR"])
        runs = doc.paragraphs[0].runs
        # Assert
        assert runs[1].bold is True

    def test_first_non_token_run_is_not_bold(self):
        """The leading non-token run should not be bold."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["foo BAR baz"])
        # Act
        preserve_bold_tokens(doc, ["BAR"])
        runs = doc.paragraphs[0].runs
        # Assert
        assert runs[0].bold is False

    def test_trailing_non_token_run_is_not_bold(self):
        """The trailing non-token run should not be bold."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["foo BAR baz"])
        # Act
        preserve_bold_tokens(doc, ["BAR"])
        runs = doc.paragraphs[0].runs
        # Assert
        assert runs[2].bold is False

    def test_token_run_text_matches_token(self):
        """The token run text should equal the matched token."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["foo BAR baz"])
        # Act
        preserve_bold_tokens(doc, ["BAR"])
        runs = doc.paragraphs[0].runs
        # Assert
        assert runs[1].text == "BAR"


class TestPreserveBoldTokensFont:
    """preserve_bold_tokens should apply the requested font to token runs."""

    def test_token_run_uses_default_ms_gothic_font(self):
        """Default font should be MS Gothic on token runs."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["foo BAR baz"])
        # Act
        preserve_bold_tokens(doc, ["BAR"])
        runs = doc.paragraphs[0].runs
        # Assert
        assert runs[1].font.name == "MS Gothic"

    def test_token_run_uses_custom_font_when_supplied(self):
        """preserve_bold_tokens should honor a custom font_name argument."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["foo BAR baz"])
        # Act
        preserve_bold_tokens(doc, ["BAR"], font_name="Yu Gothic")
        runs = doc.paragraphs[0].runs
        # Assert
        assert runs[1].font.name == "Yu Gothic"


class TestPreserveBoldTokensMultiple:
    """Multiple token occurrences and multiple tokens should both work."""

    def test_two_occurrences_of_same_token_each_become_bold_run(self):
        """Each occurrence of a token should become its own bold run."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["BAR before BAR after"])
        # Act
        preserve_bold_tokens(doc, ["BAR"])
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        # Assert
        assert len(bold_runs) == 2

    def test_multiple_tokens_in_one_paragraph_all_bolded(self):
        """All distinct tokens hit in a paragraph should be bolded."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["JST and BOOST and other stuff"])
        # Act
        preserve_bold_tokens(doc, ["JST", "BOOST"])
        runs = doc.paragraphs[0].runs
        bold_texts = sorted(r.text for r in runs if r.bold)
        # Assert
        assert bold_texts == ["BOOST", "JST"]

    def test_longer_token_wins_over_shorter_overlap_count(self):
        """If tokens overlap, exactly one bold run should be produced."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["abcdef in text"])
        # Act
        preserve_bold_tokens(doc, ["abc", "abcdef"])
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        # Assert
        assert len(bold_runs) == 1

    def test_longer_token_wins_over_shorter_overlap_text(self):
        """If tokens overlap, the longer token's text should be the one bolded."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["abcdef in text"])
        # Act
        preserve_bold_tokens(doc, ["abc", "abcdef"])
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        # Assert
        assert bold_runs[0].text == "abcdef"


class TestPreserveBoldTokensNoOp:
    """No-op cases should leave the document untouched."""

    def test_empty_token_list_is_noop(self):
        """preserve_bold_tokens with no tokens should return same doc unchanged."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["plain text here"])
        original_text = doc.paragraphs[0].text
        # Act
        preserve_bold_tokens(doc, [])
        # Assert
        assert doc.paragraphs[0].text == original_text

    def test_no_matching_tokens_leaves_paragraph_text_intact(self):
        """If no tokens match, the paragraph text should be unchanged."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["plain text here"])
        # Act
        preserve_bold_tokens(doc, ["NOT_PRESENT"])
        # Assert
        assert doc.paragraphs[0].text == "plain text here"


class TestPreserveBoldTokensCaseSensitivity:
    """case_sensitive flag controls matching."""

    def test_case_sensitive_match_misses_lowercase(self):
        """Case-sensitive match (default) should not bold lowercase variant."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["boost is lowercase"])
        # Act
        preserve_bold_tokens(doc, ["BOOST"], case_sensitive=True)
        runs = doc.paragraphs[0].runs
        # Assert
        assert all(not r.bold for r in runs)

    def test_case_insensitive_match_hits_lowercase(self):
        """Case-insensitive mode should bold a lowercase variant."""
        # Arrange
        from scitex_msword.bold import preserve_bold_tokens

        doc = _doc_with_text(["boost is lowercase"])
        # Act
        preserve_bold_tokens(doc, ["BOOST"], case_sensitive=False)
        runs = doc.paragraphs[0].runs
        # Assert
        assert any(r.bold for r in runs)


if __name__ == "__main__":
    import os

    pytest.main([os.path.abspath(__file__)])
