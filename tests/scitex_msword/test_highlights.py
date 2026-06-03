#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-02 00:00:00
# File: tests/scitex_msword/test_highlights.py

"""Tests for scitex_msword.highlights module."""

import pytest

pytestmark = pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="python-docx not installed",
)


def _doc_with_paragraphs(paragraphs):
    """Build an in-memory python-docx Document with the given run text lists."""
    import docx

    doc = docx.Document()
    for runs in paragraphs:
        p = doc.add_paragraph()
        for text in runs:
            p.add_run(text)
    return doc


class TestImportSurface:
    """Smoke tests for module imports."""

    def test_mark_additions_is_importable(self):
        """mark_additions should be importable from the module."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.highlights")
        # Assert
        assert callable(mod.mark_additions)

    def test_mark_modifications_is_importable(self):
        """mark_modifications should be importable from the module."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.highlights")
        # Assert
        assert callable(mod.mark_modifications)

    def test_extract_highlights_is_importable(self):
        """extract_highlights should be importable from the module."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.highlights")
        # Assert
        assert callable(mod.extract_highlights)


class TestMarkAdditions:
    """Tests for mark_additions."""

    def test_mark_additions_sets_highlight_color_on_target_run(self):
        """mark_additions should set highlight_color on the targeted run."""
        # Arrange
        from docx.enum.text import WD_COLOR_INDEX

        from scitex_msword.highlights import mark_additions

        doc = _doc_with_paragraphs([["hello", "world"], ["foo"]])
        # Act
        mark_additions(doc, [(0, 1)])
        # Assert
        assert doc.paragraphs[0].runs[1].font.highlight_color == WD_COLOR_INDEX.TURQUOISE

    def test_mark_additions_default_color_is_turquoise(self):
        """The default highlight color should be turquoise."""
        # Arrange
        from scitex_msword.highlights import ADDITION_COLOR

        # Act
        value = ADDITION_COLOR
        # Assert
        assert value == "turquoise"

    def test_mark_additions_accepts_custom_color(self):
        """mark_additions should accept a custom color name."""
        # Arrange
        from docx.enum.text import WD_COLOR_INDEX

        from scitex_msword.highlights import mark_additions

        doc = _doc_with_paragraphs([["hello"]])
        # Act
        mark_additions(doc, [(0, 0)], color="yellow")
        # Assert
        assert doc.paragraphs[0].runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW

    def test_mark_additions_skips_out_of_range_targets(self):
        """Out-of-range (paragraph_idx, run_idx) pairs should be silently skipped."""
        # Arrange
        from scitex_msword.highlights import mark_additions

        doc = _doc_with_paragraphs([["hello"]])
        # Act
        result = mark_additions(doc, [(99, 99), (0, 99), (99, 0)])
        # Assert
        assert result is doc

    def test_mark_additions_returns_same_document(self):
        """mark_additions should return the same Document object."""
        # Arrange
        from scitex_msword.highlights import mark_additions

        doc = _doc_with_paragraphs([["hello"]])
        # Act
        result = mark_additions(doc, [(0, 0)])
        # Assert
        assert result is doc

    def test_mark_additions_rejects_unknown_color(self):
        """An unknown color name should raise ValueError."""
        # Arrange
        from scitex_msword.highlights import mark_additions

        doc = _doc_with_paragraphs([["hello"]])
        ctx = pytest.raises(ValueError)
        # Act
        # Assert
        with ctx:
            mark_additions(doc, [(0, 0)], color="not-a-color")


class TestMarkModifications:
    """Tests for mark_modifications."""

    def test_mark_modifications_sets_pink_alias_magenta_by_default(self):
        """The default modification color should map to Word's PINK enum."""
        # Arrange
        from docx.enum.text import WD_COLOR_INDEX

        from scitex_msword.highlights import mark_modifications

        doc = _doc_with_paragraphs([["edit-me"]])
        # Act
        mark_modifications(doc, [(0, 0)])
        # Assert
        assert doc.paragraphs[0].runs[0].font.highlight_color == WD_COLOR_INDEX.PINK


class TestExtractHighlights:
    """Tests for extract_highlights."""

    def test_extract_highlights_by_color_returns_color_buckets(self):
        """extract_highlights(by_color=True) should bucket by color name."""
        # Arrange
        from scitex_msword.highlights import (
            extract_highlights,
            mark_additions,
            mark_modifications,
        )

        doc = _doc_with_paragraphs([["a"], ["b"], ["c"]])
        mark_additions(doc, [(0, 0)])
        mark_modifications(doc, [(1, 0)])
        # Act
        buckets = extract_highlights(doc, by_color=True)
        # Assert
        assert "turquoise" in buckets

    def test_extract_highlights_records_paragraph_index(self):
        """Each extracted entry should record the paragraph index."""
        # Arrange
        from scitex_msword.highlights import extract_highlights, mark_additions

        doc = _doc_with_paragraphs([["a"], ["b"]])
        mark_additions(doc, [(1, 0)])
        # Act
        buckets = extract_highlights(doc)
        # Assert
        assert buckets["turquoise"][0]["paragraph"] == 1

    def test_extract_highlights_records_run_text(self):
        """Each extracted entry should record the run text."""
        # Arrange
        from scitex_msword.highlights import extract_highlights, mark_additions

        doc = _doc_with_paragraphs([["hello world"]])
        mark_additions(doc, [(0, 0)])
        # Act
        buckets = extract_highlights(doc)
        # Assert
        assert buckets["turquoise"][0]["text"] == "hello world"

    def test_extract_highlights_no_highlights_returns_empty_dict(self):
        """If no highlights are present, the result should be empty."""
        # Arrange
        from scitex_msword.highlights import extract_highlights

        doc = _doc_with_paragraphs([["no highlight"]])
        # Act
        buckets = extract_highlights(doc)
        # Assert
        assert buckets == {}

    def test_extract_highlights_flat_mode_uses_all_bucket(self):
        """by_color=False should put all entries under 'all'."""
        # Arrange
        from scitex_msword.highlights import extract_highlights, mark_additions

        doc = _doc_with_paragraphs([["x"]])
        mark_additions(doc, [(0, 0)])
        # Act
        flat = extract_highlights(doc, by_color=False)
        # Assert
        assert "all" in flat


class TestClearHighlights:
    """Tests for clear_highlights."""

    def test_clear_highlights_removes_all_when_no_filter(self):
        """clear_highlights() with no filter should clear every highlight."""
        # Arrange
        from scitex_msword.highlights import (
            clear_highlights,
            extract_highlights,
            mark_additions,
        )

        doc = _doc_with_paragraphs([["a"], ["b"]])
        mark_additions(doc, [(0, 0), (1, 0)])
        # Act
        clear_highlights(doc)
        # Assert
        assert extract_highlights(doc) == {}

    def test_clear_highlights_filters_by_color(self):
        """clear_highlights(colors=[...]) should only clear matching colors."""
        # Arrange
        from scitex_msword.highlights import (
            clear_highlights,
            extract_highlights,
            mark_additions,
            mark_modifications,
        )

        doc = _doc_with_paragraphs([["a"], ["b"]])
        mark_additions(doc, [(0, 0)])
        mark_modifications(doc, [(1, 0)])
        # Act
        clear_highlights(doc, colors=["turquoise"])
        # Assert
        after = extract_highlights(doc)
        assert "turquoise" not in after


if __name__ == "__main__":
    import os

    pytest.main([os.path.abspath(__file__)])
