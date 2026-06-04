#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-02 00:00:00
# File: tests/scitex_msword/test_comments.py

"""Tests for scitex_msword.comments module."""

import shutil
import zipfile
from pathlib import Path

import pytest

pytestmark = pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="python-docx not installed",
)


_COMMENTS_XML = """<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0" w:author="Reviewer A" w:date="2026-06-02T10:00:00Z" w:initials="RA">
    <w:p><w:r><w:t>Please rephrase</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="1" w:author="Reviewer B" w:date="2026-06-02T11:00:00Z" w:initials="RB">
    <w:p><w:r><w:t>REPLACE: brand new wording</w:t></w:r></w:p>
  </w:comment>
</w:comments>"""

_COMMENTS_CONTENT_TYPE_OVERRIDE = (
    '<Override PartName="/word/comments.xml" '
    'ContentType="application/vnd.openxmlformats-officedocument'
    '.wordprocessingml.comments+xml"/>'
)


def _build_docx_with_comments(tmp_path: Path) -> Path:
    """
    Build a tiny docx containing two comments with anchored ranges.

    Paragraph 0: "intro paragraph"
    Paragraph 1: "anchor one text here" — comment id=0 wraps "anchor one"
    Paragraph 2: "anchor two text here" — comment id=1 wraps "anchor two"
    """
    import docx

    base = tmp_path / "base.docx"
    doc = docx.Document()
    doc.add_paragraph("intro paragraph")
    doc.add_paragraph("anchor one text here")
    doc.add_paragraph("anchor two text here")
    doc.save(str(base))

    # Now inject commentRangeStart/End markers + a wired comments.xml.
    out = tmp_path / "with_comments.docx"
    shutil.copy(base, out)

    with zipfile.ZipFile(out, "r") as zin:
        members = {name: zin.read(name) for name in zin.namelist()}

    document_xml = members["word/document.xml"].decode("utf-8")

    # Inject comment ranges around the literal phrases.
    def wrap(xml: str, phrase: str, cid: str) -> str:
        target = f"<w:t>{phrase}"
        if target not in xml:
            return xml
        replacement = (
            f'<w:commentRangeStart w:id="{cid}"/>'
            f"<w:t>{phrase}"
        )
        xml = xml.replace(target, replacement, 1)
        # Close: insert before the next </w:t>... actually easier: insert
        # commentRangeEnd just before </w:p> of that paragraph.
        # We rely on each phrase being unique per paragraph.
        return xml

    document_xml = wrap(document_xml, "anchor one", "0")
    document_xml = wrap(document_xml, "anchor two", "1")

    # Close ranges at end of each paragraph containing the marker.
    # We do a targeted close by injecting commentRangeEnd right before
    # the closing </w:p> that follows the matching commentRangeStart.
    def close_after(xml: str, cid: str) -> str:
        marker = f'<w:commentRangeStart w:id="{cid}"/>'
        idx = xml.find(marker)
        if idx < 0:
            return xml
        end_p = xml.find("</w:p>", idx)
        if end_p < 0:
            return xml
        insertion = f'<w:commentRangeEnd w:id="{cid}"/>'
        return xml[:end_p] + insertion + xml[end_p:]

    document_xml = close_after(document_xml, "0")
    document_xml = close_after(document_xml, "1")

    members["word/document.xml"] = document_xml.encode("utf-8")
    members["word/comments.xml"] = _COMMENTS_XML.encode("utf-8")

    # Append a comments Override into the existing [Content_Types].xml
    # (preserving the defaults for jpeg/png/rels/xml that python-docx wrote).
    ct = members["[Content_Types].xml"].decode("utf-8")
    if "comments.xml" not in ct:
        ct = ct.replace(
            "</Types>", _COMMENTS_CONTENT_TYPE_OVERRIDE + "</Types>", 1
        )
        members["[Content_Types].xml"] = ct.encode("utf-8")

    # Append a relationship for comments.xml.
    rels_key = "word/_rels/document.xml.rels"
    rels = members[rels_key].decode("utf-8")
    if "comments.xml" not in rels:
        rels = rels.replace(
            "</Relationships>",
            '<Relationship Id="rIdComments" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
            'Target="comments.xml"/></Relationships>',
            1,
        )
        members[rels_key] = rels.encode("utf-8")

    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in members.items():
            zout.writestr(name, data)

    return out


class TestImportSurface:
    """Smoke tests."""

    def test_extract_comments_is_importable(self):
        """extract_comments should import cleanly."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.comments")
        # Assert
        assert callable(mod.extract_comments)

    def test_apply_comments_as_edits_is_importable(self):
        """apply_comments_as_edits should import cleanly."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.comments")
        # Assert
        assert callable(mod.apply_comments_as_edits)


class TestExtractCommentsBasic:
    """extract_comments should parse comments.xml."""

    def test_extract_comments_returns_two_entries(self, tmp_path):
        """The fixture has two comments; both should be returned."""
        # Arrange
        from scitex_msword.comments import extract_comments

        path = _build_docx_with_comments(tmp_path)
        # Act
        comments = extract_comments(path)
        # Assert
        assert len(comments) == 2

    def test_extract_comments_records_author(self, tmp_path):
        """The first comment's author should be 'Reviewer A'."""
        # Arrange
        from scitex_msword.comments import extract_comments

        path = _build_docx_with_comments(tmp_path)
        # Act
        comments = extract_comments(path)
        # Assert
        assert comments[0]["author"] == "Reviewer A"

    def test_extract_comments_records_text(self, tmp_path):
        """Each comment's body text should be extracted."""
        # Arrange
        from scitex_msword.comments import extract_comments

        path = _build_docx_with_comments(tmp_path)
        # Act
        comments = extract_comments(path)
        bodies = sorted(c["text"] for c in comments)
        # Assert
        assert bodies == ["Please rephrase", "REPLACE: brand new wording"]


class TestExtractCommentsNoComments:
    """When no comments.xml part exists, the result should be an empty list."""

    def test_docx_without_comments_returns_empty_list(self, tmp_path):
        """A vanilla docx should have no comments."""
        # Arrange
        import docx

        from scitex_msword.comments import extract_comments

        path = tmp_path / "plain.docx"
        doc = docx.Document()
        doc.add_paragraph("just a line")
        doc.save(str(path))
        # Act
        comments = extract_comments(path)
        # Assert
        assert comments == []


class TestExtractCommentsAnchors:
    """When commentRangeStart/End are present, anchor info should populate."""

    def test_anchor_text_captures_anchor_phrase(self, tmp_path):
        """The anchor_text for comment 0 should include 'anchor one'."""
        # Arrange
        from scitex_msword.comments import extract_comments

        path = _build_docx_with_comments(tmp_path)
        # Act
        comments = extract_comments(path)
        c0 = next(c for c in comments if c["id"] == 0)
        # Assert
        assert "anchor one" in c0["anchor_text"]

    def test_paragraph_range_for_anchor_zero_is_paragraph_one(self, tmp_path):
        """Comment 0 anchors paragraph index 1."""
        # Arrange
        from scitex_msword.comments import extract_comments

        path = _build_docx_with_comments(tmp_path)
        # Act
        comments = extract_comments(path)
        c0 = next(c for c in comments if c["id"] == 0)
        # Assert
        assert c0["paragraph_range"][0] == 1


class TestApplyCommentsAsEdits:
    """apply_comments_as_edits should honor the REPLACE: grammar."""

    def test_apply_comments_returns_summary_dict(self, tmp_path):
        """The summary dict should include applied/skipped counts."""
        # Arrange
        import docx

        from scitex_msword.comments import apply_comments_as_edits, extract_comments

        path = _build_docx_with_comments(tmp_path)
        doc = docx.Document(str(path))
        # Act
        # Use pre-extracted comments (extract_comments needs the path/zip).
        comments = extract_comments(path)
        summary = apply_comments_as_edits(doc, comments=comments)
        # Assert
        assert "applied" in summary and "skipped" in summary

    def test_apply_comments_applies_replace_grammar(self, tmp_path):
        """A REPLACE: comment with an anchor should perform the substitution."""
        # Arrange
        import docx

        from scitex_msword.comments import apply_comments_as_edits, extract_comments

        path = _build_docx_with_comments(tmp_path)
        doc = docx.Document(str(path))
        comments = extract_comments(path)
        # Act
        summary = apply_comments_as_edits(doc, comments=comments)
        # Assert: at least one REPLACE comment was applied.
        assert summary["applied"] >= 1

    def test_apply_comments_skips_non_replace_comments(self, tmp_path):
        """Plain (non-REPLACE) comments should be reported as skipped."""
        # Arrange
        import docx

        from scitex_msword.comments import apply_comments_as_edits, extract_comments

        path = _build_docx_with_comments(tmp_path)
        doc = docx.Document(str(path))
        comments = extract_comments(path)
        # Act
        summary = apply_comments_as_edits(doc, comments=comments)
        # Assert
        assert summary["skipped"] >= 1

    def test_apply_comments_rejects_unknown_grammar(self, tmp_path):
        """A grammar argument other than 'replace' should raise ValueError."""
        # Arrange
        import docx

        from scitex_msword.comments import apply_comments_as_edits

        path = _build_docx_with_comments(tmp_path)
        doc = docx.Document(str(path))
        ctx = pytest.raises(ValueError)
        # Act
        # Assert
        with ctx:
            apply_comments_as_edits(doc, comments=[], grammar="natural-language")


if __name__ == "__main__":
    import os

    pytest.main([os.path.abspath(__file__)])
