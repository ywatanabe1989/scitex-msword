#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-02 00:00:00
# File: tests/scitex_msword/test_track_changes.py

"""Tests for scitex_msword.track_changes module."""

import io

import pytest

pytestmark = pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="python-docx not installed",
)


def _doc_with_paragraphs(paragraphs):
    """Build an in-memory python-docx Document with the given run text lists."""
    import docx

    doc = docx.Document()
    for runs in paragraphs:
        p = doc.add_paragraph()
        for text in runs:
            p.add_run(text)
    return doc


def _roundtrip(doc):
    """Save to BytesIO then re-load, simulating disk round-trip."""
    import docx

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return docx.Document(buf)


# ---------------------------------------------------------------------------
# Import surface
# ---------------------------------------------------------------------------


class TestImportSurface:
    """Smoke tests for module imports."""

    def test_enable_track_changes_is_importable(self):
        """enable_track_changes should be importable from the module."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.track_changes")
        # Assert
        assert callable(mod.enable_track_changes)

    def test_wrap_insertion_is_importable(self):
        """wrap_as_tracked_insertion should be importable from the module."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.track_changes")
        # Assert
        assert callable(mod.wrap_as_tracked_insertion)

    def test_wrap_deletion_is_importable(self):
        """wrap_as_tracked_deletion should be importable from the module."""
        # Arrange
        import importlib

        # Act
        mod = importlib.import_module("scitex_msword.track_changes")
        # Assert
        assert callable(mod.wrap_as_tracked_deletion)

    def test_top_level_reexports(self):
        """Top-level scitex_msword should re-export the track-changes API."""
        # Arrange
        import scitex_msword as sxm

        # Act
        names = (
            "enable_track_changes",
            "wrap_as_tracked_insertion",
            "wrap_as_tracked_deletion",
            "extract_tracked_changes",
            "accept_all_tracked_changes",
            "reject_all_tracked_changes",
        )
        # Assert
        for name in names:
            assert callable(getattr(sxm, name)), name


# ---------------------------------------------------------------------------
# enable_track_changes
# ---------------------------------------------------------------------------


class TestEnableTrackChanges:
    """Tests for enable_track_changes."""

    def test_enable_inserts_trackChanges_element(self):
        """enable_track_changes(True) should insert <w:trackChanges/>."""
        # Arrange
        from scitex_msword.track_changes import (
            enable_track_changes,
            is_track_changes_enabled,
        )

        doc = _doc_with_paragraphs([["hi"]])
        # Act
        enable_track_changes(doc, enabled=True)
        # Assert
        assert is_track_changes_enabled(doc)

    def test_enable_persists_through_save_and_reload(self):
        """The flag should survive a save -> reload round-trip."""
        # Arrange
        from scitex_msword.track_changes import (
            enable_track_changes,
            is_track_changes_enabled,
        )

        doc = _doc_with_paragraphs([["hi"]])
        # Act
        enable_track_changes(doc)
        reloaded = _roundtrip(doc)
        # Assert
        assert is_track_changes_enabled(reloaded)

    def test_enable_is_idempotent(self):
        """Calling enable_track_changes twice should yield exactly one element."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import enable_track_changes

        doc = _doc_with_paragraphs([["hi"]])
        # Act
        enable_track_changes(doc)
        enable_track_changes(doc)
        elems = doc.settings.element.findall(qn("w:trackChanges"))
        # Assert
        assert len(elems) == 1

    def test_disable_removes_trackChanges_element(self):
        """enable_track_changes(False) should remove existing trackChanges."""
        # Arrange
        from scitex_msword.track_changes import (
            enable_track_changes,
            is_track_changes_enabled,
        )

        doc = _doc_with_paragraphs([["hi"]])
        enable_track_changes(doc, enabled=True)
        # Act
        enable_track_changes(doc, enabled=False)
        # Assert
        assert not is_track_changes_enabled(doc)

    def test_disable_when_absent_is_noop(self):
        """Disabling when no trackChanges is present should not raise."""
        # Arrange
        from scitex_msword.track_changes import (
            enable_track_changes,
            is_track_changes_enabled,
        )

        doc = _doc_with_paragraphs([["hi"]])
        # Act
        enable_track_changes(doc, enabled=False)
        # Assert
        assert not is_track_changes_enabled(doc)

    def test_enable_returns_same_document(self):
        """enable_track_changes should return the same Document (chainable)."""
        # Arrange
        from scitex_msword.track_changes import enable_track_changes

        doc = _doc_with_paragraphs([["hi"]])
        # Act
        result = enable_track_changes(doc)
        # Assert
        assert result is doc


# ---------------------------------------------------------------------------
# wrap_as_tracked_insertion
# ---------------------------------------------------------------------------


class TestWrapAsTrackedInsertion:
    """Tests for wrap_as_tracked_insertion."""

    def test_wrap_creates_w_ins_element(self):
        """A wrapped run should now sit inside a <w:ins> parent."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_insertion

        doc = _doc_with_paragraphs([["a", "b"]])
        p = doc.paragraphs[0]
        target_run = p.runs[1]
        # Act
        wrap_as_tracked_insertion(p, runs=[1], author="agent")
        # Assert
        assert target_run._r.getparent().tag == qn("w:ins")

    def test_wrap_sets_author_attribute(self):
        """The <w:ins> wrapper should carry the author attribute."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_insertion

        doc = _doc_with_paragraphs([["x", "y"]])
        p = doc.paragraphs[0]
        # Act
        wraps = wrap_as_tracked_insertion(p, runs=[0], author="reviewer")
        # Assert
        assert wraps[0].get(qn("w:author")) == "reviewer"

    def test_wrap_sets_explicit_id(self):
        """When w_id is supplied it should be honored verbatim."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_insertion

        doc = _doc_with_paragraphs([["x"]])
        p = doc.paragraphs[0]
        # Act
        wraps = wrap_as_tracked_insertion(p, runs=[0], w_id=42)
        # Assert
        assert wraps[0].get(qn("w:id")) == "42"

    def test_wrap_default_date_is_iso(self):
        """The default <w:date> should be an ISO-8601 UTC string."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_insertion

        doc = _doc_with_paragraphs([["x"]])
        p = doc.paragraphs[0]
        # Act
        wraps = wrap_as_tracked_insertion(p, runs=[0])
        # Assert
        assert wraps[0].get(qn("w:date")).endswith("Z")

    def test_wrap_accepts_run_object(self):
        """A Run object (not just an index) should be accepted as a target."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_insertion

        doc = _doc_with_paragraphs([["x", "y"]])
        p = doc.paragraphs[0]
        target = p.runs[0]
        # Act
        wrap_as_tracked_insertion(p, runs=[target])
        # Assert
        assert target._r.getparent().tag == qn("w:ins")

    def test_wrap_skips_unknown_indices(self):
        """Out-of-range indices should be silently skipped (no wrappers made)."""
        # Arrange
        from scitex_msword.track_changes import wrap_as_tracked_insertion

        doc = _doc_with_paragraphs([["x"]])
        p = doc.paragraphs[0]
        # Act
        wraps = wrap_as_tracked_insertion(p, runs=[99])
        # Assert
        assert wraps == []

    def test_wrap_auto_id_increments(self):
        """Two successive wraps should produce strictly increasing ids."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_insertion

        doc = _doc_with_paragraphs([["a", "b"]])
        p = doc.paragraphs[0]
        # Act
        w1 = wrap_as_tracked_insertion(p, runs=[0])
        w2 = wrap_as_tracked_insertion(p, runs=[0])  # 0 again after re-eval
        # Assert
        id1 = int(w1[0].get(qn("w:id")))
        id2 = int(w2[0].get(qn("w:id")))
        assert id2 > id1


# ---------------------------------------------------------------------------
# wrap_as_tracked_deletion
# ---------------------------------------------------------------------------


class TestWrapAsTrackedDeletion:
    """Tests for wrap_as_tracked_deletion."""

    def test_wrap_creates_w_del_element(self):
        """A wrapped run should now sit inside a <w:del> parent."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_deletion

        doc = _doc_with_paragraphs([["a", "b"]])
        p = doc.paragraphs[0]
        target_run = p.runs[1]
        # Act
        wrap_as_tracked_deletion(p, runs=[1], author="agent")
        # Assert
        assert target_run._r.getparent().tag == qn("w:del")

    def test_wrap_converts_t_to_delText(self):
        """The inner <w:t> children should be retagged as <w:delText>."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_deletion

        doc = _doc_with_paragraphs([["delete-me"]])
        p = doc.paragraphs[0]
        # Act
        wraps = wrap_as_tracked_deletion(p, runs=[0])
        # Assert
        assert len(list(wraps[0].iter(qn("w:delText")))) == 1

    def test_wrap_removes_original_t_tag(self):
        """No <w:t> elements should remain inside the deletion wrapper."""
        # Arrange
        from docx.oxml.ns import qn

        from scitex_msword.track_changes import wrap_as_tracked_deletion

        doc = _doc_with_paragraphs([["delete-me"]])
        p = doc.paragraphs[0]
        # Act
        wraps = wrap_as_tracked_deletion(p, runs=[0])
        # Assert
        assert list(wraps[0].iter(qn("w:t"))) == []


# ---------------------------------------------------------------------------
# extract_tracked_changes
# ---------------------------------------------------------------------------


class TestExtractTrackedChanges:
    """Tests for extract_tracked_changes."""

    def test_extract_returns_insertions(self):
        """extract_tracked_changes should surface inserted runs."""
        # Arrange
        from scitex_msword.track_changes import (
            extract_tracked_changes,
            wrap_as_tracked_insertion,
        )

        doc = _doc_with_paragraphs([["alpha", "beta"]])
        wrap_as_tracked_insertion(doc.paragraphs[0], runs=[1], author="a")
        # Act
        changes = extract_tracked_changes(doc)
        # Assert
        assert [c["type"] for c in changes] == ["insert"]

    def test_extract_returns_deletions(self):
        """extract_tracked_changes should surface deleted runs."""
        # Arrange
        from scitex_msword.track_changes import (
            extract_tracked_changes,
            wrap_as_tracked_deletion,
        )

        doc = _doc_with_paragraphs([["alpha", "beta"]])
        wrap_as_tracked_deletion(doc.paragraphs[0], runs=[1], author="a")
        # Act
        changes = extract_tracked_changes(doc)
        # Assert
        assert [c["type"] for c in changes] == ["delete"]

    def test_extract_records_text_content(self):
        """Each entry should record the wrapped text."""
        # Arrange
        from scitex_msword.track_changes import (
            extract_tracked_changes,
            wrap_as_tracked_insertion,
        )

        doc = _doc_with_paragraphs([["alpha", "beta"]])
        wrap_as_tracked_insertion(doc.paragraphs[0], runs=[0])
        # Act
        changes = extract_tracked_changes(doc)
        # Assert
        assert changes[0]["text"] == "alpha"

    def test_extract_records_paragraph_index(self):
        """The paragraph_idx field should be populated."""
        # Arrange
        from scitex_msword.track_changes import (
            extract_tracked_changes,
            wrap_as_tracked_insertion,
        )

        doc = _doc_with_paragraphs([["a"], ["b"], ["c"]])
        wrap_as_tracked_insertion(doc.paragraphs[2], runs=[0])
        # Act
        changes = extract_tracked_changes(doc)
        # Assert
        assert changes[0]["paragraph_idx"] == 2

    def test_extract_empty_when_no_revisions(self):
        """A pristine document should yield an empty list."""
        # Arrange
        from scitex_msword.track_changes import extract_tracked_changes

        doc = _doc_with_paragraphs([["plain"]])
        # Act
        changes = extract_tracked_changes(doc)
        # Assert
        assert changes == []


# ---------------------------------------------------------------------------
# accept_all / reject_all
# ---------------------------------------------------------------------------


class TestAcceptRejectAll:
    """Tests for accept_all_tracked_changes / reject_all_tracked_changes."""

    def test_accept_unwraps_insertions_no_revisions_left(self):
        """Accepting should leave no tracked changes."""
        # Arrange
        from scitex_msword.track_changes import (
            accept_all_tracked_changes,
            extract_tracked_changes,
            wrap_as_tracked_insertion,
        )

        doc = _doc_with_paragraphs([["keep-me"]])
        wrap_as_tracked_insertion(doc.paragraphs[0], runs=[0])
        # Act
        accept_all_tracked_changes(doc)
        # Assert
        assert extract_tracked_changes(doc) == []

    def test_accept_unwraps_insertions_preserves_text(self):
        """Accepting an insertion should keep the inserted content."""
        # Arrange
        from scitex_msword.track_changes import (
            accept_all_tracked_changes,
            wrap_as_tracked_insertion,
        )

        doc = _doc_with_paragraphs([["keep-me"]])
        wrap_as_tracked_insertion(doc.paragraphs[0], runs=[0])
        # Act
        accept_all_tracked_changes(doc)
        # Assert
        assert doc.paragraphs[0].text == "keep-me"

    def test_accept_drops_deletions_no_revisions_left(self):
        """Accepting should leave no tracked changes from deletions."""
        # Arrange
        from scitex_msword.track_changes import (
            accept_all_tracked_changes,
            extract_tracked_changes,
            wrap_as_tracked_deletion,
        )

        doc = _doc_with_paragraphs([["keep ", "drop"]])
        wrap_as_tracked_deletion(doc.paragraphs[0], runs=[1])
        # Act
        accept_all_tracked_changes(doc)
        # Assert
        assert extract_tracked_changes(doc) == []

    def test_accept_drops_deletions_removes_text(self):
        """Accepting a deletion should drop the deleted content."""
        # Arrange
        from scitex_msword.track_changes import (
            accept_all_tracked_changes,
            wrap_as_tracked_deletion,
        )

        doc = _doc_with_paragraphs([["keep ", "drop"]])
        wrap_as_tracked_deletion(doc.paragraphs[0], runs=[1])
        # Act
        accept_all_tracked_changes(doc)
        # Assert
        assert doc.paragraphs[0].text.strip() == "keep"

    def test_reject_drops_insertions_no_revisions_left(self):
        """Rejecting should leave no tracked changes from insertions."""
        # Arrange
        from scitex_msword.track_changes import (
            extract_tracked_changes,
            reject_all_tracked_changes,
            wrap_as_tracked_insertion,
        )

        doc = _doc_with_paragraphs([["keep ", "drop"]])
        wrap_as_tracked_insertion(doc.paragraphs[0], runs=[1])
        # Act
        reject_all_tracked_changes(doc)
        # Assert
        assert extract_tracked_changes(doc) == []

    def test_reject_drops_insertions_removes_text(self):
        """Rejecting an insertion should drop the inserted content."""
        # Arrange
        from scitex_msword.track_changes import (
            reject_all_tracked_changes,
            wrap_as_tracked_insertion,
        )

        doc = _doc_with_paragraphs([["keep ", "drop"]])
        wrap_as_tracked_insertion(doc.paragraphs[0], runs=[1])
        # Act
        reject_all_tracked_changes(doc)
        # Assert
        assert doc.paragraphs[0].text.strip() == "keep"

    def test_reject_restores_deletions_no_revisions_left(self):
        """Rejecting a deletion should leave no tracked changes."""
        # Arrange
        from scitex_msword.track_changes import (
            extract_tracked_changes,
            reject_all_tracked_changes,
            wrap_as_tracked_deletion,
        )

        doc = _doc_with_paragraphs([["restored-text"]])
        wrap_as_tracked_deletion(doc.paragraphs[0], runs=[0])
        # Act
        reject_all_tracked_changes(doc)
        # Assert
        assert extract_tracked_changes(doc) == []

    def test_reject_restores_deletions_as_text(self):
        """Rejecting a deletion should restore the original text content."""
        # Arrange
        from scitex_msword.track_changes import (
            reject_all_tracked_changes,
            wrap_as_tracked_deletion,
        )

        doc = _doc_with_paragraphs([["restored-text"]])
        wrap_as_tracked_deletion(doc.paragraphs[0], runs=[0])
        # Act
        reject_all_tracked_changes(doc)
        # Assert
        assert "restored-text" in doc.paragraphs[0].text


# ---------------------------------------------------------------------------
# Module-level smoke through save / reload
# ---------------------------------------------------------------------------


class TestRoundTrip:
    """End-to-end smoke through save -> reload."""

    def test_insertion_survives_save_and_reload(self):
        """A wrapped insertion should be readable from the reloaded file."""
        # Arrange
        from scitex_msword.track_changes import (
            extract_tracked_changes,
            wrap_as_tracked_insertion,
        )

        doc = _doc_with_paragraphs([["alpha", "beta"]])
        wrap_as_tracked_insertion(doc.paragraphs[0], runs=[1], author="rt")
        # Act
        reloaded = _roundtrip(doc)
        changes = extract_tracked_changes(reloaded)
        # Assert
        assert changes and changes[0]["text"] == "beta"


if __name__ == "__main__":
    import os

    pytest.main([os.path.abspath(__file__)])
