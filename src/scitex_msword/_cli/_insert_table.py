#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-04 12:00:00
# File: src/scitex_msword/_cli/_insert_table.py
#
# Part of scitex-msword (AGPL-3.0-only). See LICENSE at the repo root.

"""``scitex-msword insert-table`` — Click wiring for the table-insert verb."""

from __future__ import annotations

import json as _json
from typing import List, Optional

import click


def _parse_rows_arg(rows: Optional[str], rows_file: Optional[str]) -> List[List[str]]:
    """Decode the ``--rows`` / ``--rows-file`` input into list-of-lists-of-str."""
    if rows_file is not None:
        from pathlib import Path as _Path

        raw = _Path(rows_file).read_text(encoding="utf-8")
    else:
        raw = rows or ""
    try:
        parsed = _json.loads(raw)
    except _json.JSONDecodeError as e:
        raise click.UsageError(
            f"insert-table: invalid JSON for --rows / --rows-file: {e}"
        ) from e
    if not isinstance(parsed, list) or not all(isinstance(r, list) for r in parsed):
        raise click.UsageError(
            "insert-table: --rows / --rows-file must decode to a JSON "
            "array-of-arrays-of-strings"
        )
    return [[str(c) for c in r] for r in parsed]


def _parse_col_widths_arg(col_widths_dxa: str) -> List[int]:
    """Decode the comma-separated ``--col-widths-dxa`` arg into ``List[int]``."""
    parts = [p.strip() for p in col_widths_dxa.split(",") if p.strip()]
    if not parts:
        raise click.UsageError("insert-table: --col-widths-dxa cannot be empty")
    try:
        return [int(p) for p in parts]
    except ValueError as e:
        raise click.UsageError(
            f"insert-table: --col-widths-dxa entries must be integers ({e})"
        ) from e


def register(main_group: click.Group) -> None:
    """Attach ``insert-table`` to the top-level CLI group."""

    @main_group.command("insert-table")
    @click.option("--path", required=True, type=click.Path(), help="Input .docx path.")
    @click.option("--out", required=True, type=click.Path(), help="Output .docx path.")
    @click.option(
        "--paragraph-index",
        type=int,
        required=True,
        help="0-based index of the anchor paragraph in doc.paragraphs.",
    )
    @click.option(
        "--rows",
        default=None,
        help=(
            "JSON array-of-arrays of strings (e.g. "
            '\'[["役割","モジュール"],["論文執筆","scitex-writer"]]\'). '
            "Mutually exclusive with --rows-file."
        ),
    )
    @click.option(
        "--rows-file",
        default=None,
        type=click.Path(),
        help="Path to a JSON file containing the array-of-arrays of strings.",
    )
    @click.option(
        "--col-widths-dxa",
        default="3000,6000",
        help=(
            "Comma-separated column widths in dxa (twentieths of a point). "
            "Default '3000,6000' (2-col 1:2)."
        ),
    )
    @click.option(
        "--header-row/--no-header-row",
        default=True,
        help="Style row 0 as a header (default: header on).",
    )
    @click.option(
        "--body-font", default="MS 明朝", help="Body row font (default 'MS 明朝')."
    )
    @click.option(
        "--header-font",
        default="MS ゴシック",
        help="Header row font (default 'MS ゴシック').",
    )
    @click.option(
        "--font-size-pt",
        type=float,
        default=10.5,
        help="Run font size in points (default 10.5).",
    )
    @click.option(
        "--track-changes/--no-track-changes",
        "track_changes",
        default=None,
        help=(
            "Force row-level <w:trPr><w:ins/></w:trPr> markers on "
            "(--track-changes) or off (--no-track-changes). Omit both to "
            "auto-detect from settings.xml."
        ),
    )
    @click.option(
        "--track-changes-author",
        default="agent",
        help="Author string for the row-ins markers (default 'agent').",
    )
    @click.option(
        "--track-changes-date",
        default=None,
        help="ISO-8601 timestamp for the row-ins markers (default: now UTC).",
    )
    @click.option(
        "--json", "as_json", is_flag=True, default=False, help="Output as JSON."
    )
    def insert_table_cmd(
        path,
        out,
        paragraph_index,
        rows,
        rows_file,
        col_widths_dxa,
        header_row,
        body_font,
        header_font,
        font_size_pt,
        track_changes,
        track_changes_author,
        track_changes_date,
        as_json,
    ):
        """Insert a Word table after the paragraph at PARAGRAPH_INDEX.

        Builds the ``<w:tbl>`` directly via lxml and inserts as the next
        sibling of ``doc.paragraphs[paragraph_index]``. When the source
        document has Track Changes enabled, each generated ``<w:tr>``
        is marked with ``<w:trPr><w:ins/></w:trPr>`` so Word surfaces
        the rows as accept/reject-able revisions.

        \b
        Example:
            $ scitex-msword insert-table \\
                --path draft.docx --out draft_with_table.docx \\
                --paragraph-index 17 \\
                --rows '[["役割","モジュール名"],["論文執筆","scitex-writer"]]' \\
                --col-widths-dxa 3000,6000
            $ scitex-msword insert-table \\
                --path draft.docx --out draft_with_table.docx \\
                --paragraph-index 17 --rows-file table.json \\
                --track-changes-author claude-agent
        """
        if (rows is None) == (rows_file is None):
            raise click.UsageError(
                "insert-table: provide exactly one of --rows or --rows-file"
            )
        parsed_rows = _parse_rows_arg(rows, rows_file)
        col_widths = _parse_col_widths_arg(col_widths_dxa)

        import docx as _docx
        from pathlib import Path as _Path

        from ..tables import insert_table_after_paragraph

        doc = _docx.Document(path)
        insert_table_after_paragraph(
            doc,
            paragraph_index=paragraph_index,
            rows=parsed_rows,
            col_widths_dxa=col_widths,
            header_row=header_row,
            body_font=body_font,
            header_font=header_font,
            font_size_pt=font_size_pt,
            track_changes=track_changes,
            track_changes_author=track_changes_author,
            track_changes_date=track_changes_date,
        )
        _Path(out).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)

        if as_json:
            click.echo(_json.dumps({"out": out, "rows": len(parsed_rows)}))
        else:
            click.echo(out)


__all__ = ["register"]
