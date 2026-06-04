#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-04 12:00:00
# File: src/scitex_msword/tables.py
#
# Part of scitex-msword (AGPL-3.0-only). See LICENSE at the repo root.

"""
Table insertion utilities for python-docx Documents.

The public entry point :func:`insert_table_after_paragraph` builds a
``<w:tbl>`` element directly with pure lxml (so it works on raw
python-docx 1.1.x without going through the high-level Table API) and
inserts it as the next sibling of a target paragraph. When the source
document already has Track Changes enabled, each generated ``<w:tr>``
is marked as an inserted row by emitting
``<w:trPr><w:ins .../></w:trPr>`` so that Word surfaces the new rows
as accept/reject-able revisions — matching the dogfood-time pattern
proj-grant's ``build_v43.py`` shipped in BOOST today.

OOXML refs:

- ``w:tbl``        ECMA-376 §17.4.38
- ``w:tblPr``      ECMA-376 §17.4.59
- ``w:tblGrid``    ECMA-376 §17.4.49
- ``w:tr``         ECMA-376 §17.4.79
- ``w:trPr``       ECMA-376 §17.4.82
- ``w:ins`` (in trPr)  ECMA-376 §17.13.5.18 (row-level insertion marker)
"""

from __future__ import annotations

from copy import deepcopy
from typing import Any, Iterable, List, Optional, Sequence, Tuple, Union

try:
    from docx.document import Document as DocxDocument  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]
    from lxml import etree

    DOCX_AVAILABLE = True
    _DOCX_IMPORT_ERROR: Optional[Exception] = None
except ImportError as exc:  # pragma: no cover — optional dep
    DOCX_AVAILABLE = False
    _DOCX_IMPORT_ERROR = exc
    DocxDocument = None  # type: ignore[assignment,misc]
    qn = None  # type: ignore[assignment]
    etree = None  # type: ignore[assignment]


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_NS = "http://www.w3.org/XML/1998/namespace"


def _ensure_docx_available() -> None:
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx (and lxml) are required for scitex_msword.tables. "
            "Install via `pip install python-docx`."
        ) from _DOCX_IMPORT_ERROR


def _w(tag: str) -> str:
    """Return the Clark-form qualified WordprocessingML tag for ``tag``."""
    return f"{{{_W_NS}}}{tag}"


def _font_size_to_half_points(font_size_pt: float) -> str:
    """Word's ``<w:sz w:val=…>`` uses half-points (e.g. 10.5pt → "21")."""
    return str(int(round(font_size_pt * 2)))


def _normalise_rows(rows: Iterable[Sequence[str]]) -> List[List[str]]:
    """Coerce ``rows`` into ``List[List[str]]`` with consistent column counts."""
    normalised: List[List[str]] = []
    for r in rows:
        cells = [str(c) for c in r]
        normalised.append(cells)
    if not normalised:
        raise ValueError("rows must contain at least one row")
    n_cols = len(normalised[0])
    if n_cols == 0:
        raise ValueError("rows must have at least one column")
    for i, r in enumerate(normalised):
        if len(r) != n_cols:
            raise ValueError(
                f"all rows must have the same number of columns; "
                f"row 0 has {n_cols} but row {i} has {len(r)}"
            )
    return normalised


def _build_cell(
    *,
    text: str,
    width_dxa: int,
    font_name: str,
    font_sz_half_points: str,
    bold: bool,
) -> "etree._Element":
    """Build a single ``<w:tc>`` element."""
    tc = etree.Element(_w("tc"))
    tcPr = etree.SubElement(tc, _w("tcPr"))
    tcW = etree.SubElement(tcPr, _w("tcW"))
    tcW.set(_w("w"), str(width_dxa))
    tcW.set(_w("type"), "dxa")
    vAlign = etree.SubElement(tcPr, _w("vAlign"))
    vAlign.set(_w("val"), "center")

    p = etree.SubElement(tc, _w("p"))
    r = etree.SubElement(p, _w("r"))
    rPr = etree.SubElement(r, _w("rPr"))
    rFonts = etree.SubElement(rPr, _w("rFonts"))
    for slot in ("ascii", "eastAsia", "hAnsi", "cs"):
        rFonts.set(_w(slot), font_name)
    sz = etree.SubElement(rPr, _w("sz"))
    sz.set(_w("val"), font_sz_half_points)
    szCs = etree.SubElement(rPr, _w("szCs"))
    szCs.set(_w("val"), font_sz_half_points)
    if bold:
        etree.SubElement(rPr, _w("b"))
        etree.SubElement(rPr, _w("bCs"))

    t = etree.SubElement(r, _w("t"))
    if text != text.strip() or " " in text:
        t.set(f"{{{_XML_NS}}}space", "preserve")
    t.text = text
    return tc


def _build_row(
    *,
    cells: Sequence[str],
    col_widths_dxa: Sequence[int],
    font_name: str,
    font_sz_half_points: str,
    bold: bool,
) -> "etree._Element":
    """Build a single ``<w:tr>`` element with its constituent cells."""
    tr = etree.Element(_w("tr"))
    for cell_text, width in zip(cells, col_widths_dxa):
        tr.append(
            _build_cell(
                text=cell_text,
                width_dxa=width,
                font_name=font_name,
                font_sz_half_points=font_sz_half_points,
                bold=bold,
            )
        )
    return tr


def _mark_row_as_tracked_insertion(
    tr_element: "etree._Element",
    *,
    w_id: int,
    author: str,
    date: str,
) -> None:
    """Insert ``<w:trPr><w:ins .../></w:trPr>`` so Word treats the row as inserted."""
    trPr = tr_element.find(_w("trPr"))
    if trPr is None:
        trPr = etree.Element(_w("trPr"))
        tr_element.insert(0, trPr)
    ins_el = etree.SubElement(trPr, _w("ins"))
    ins_el.set(_w("id"), str(w_id))
    ins_el.set(_w("author"), author)
    ins_el.set(_w("date"), date)


def _build_table(
    *,
    rows: Sequence[Sequence[str]],
    col_widths_dxa: Sequence[int],
    header_row: bool,
    body_font: str,
    header_font: str,
    font_sz_half_points: str,
) -> "etree._Element":
    """Build the full ``<w:tbl>`` element (no track-changes wrapping)."""
    tbl_el = etree.Element(_w("tbl"))

    # tblPr — width + borders
    tblPr = etree.SubElement(tbl_el, _w("tblPr"))
    tblW = etree.SubElement(tblPr, _w("tblW"))
    tblW.set(_w("w"), "5000")
    tblW.set(_w("type"), "pct")
    tblBorders = etree.SubElement(tblPr, _w("tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = etree.SubElement(tblBorders, _w(side))
        b.set(_w("val"), "single")
        b.set(_w("sz"), "4")
        b.set(_w("space"), "0")
        b.set(_w("color"), "auto")

    # tblGrid — per-column declared widths
    tblGrid = etree.SubElement(tbl_el, _w("tblGrid"))
    for width_dxa in col_widths_dxa:
        gc = etree.SubElement(tblGrid, _w("gridCol"))
        gc.set(_w("w"), str(width_dxa))

    # Rows
    for row_idx, cells in enumerate(rows):
        is_header = header_row and row_idx == 0
        tbl_el.append(
            _build_row(
                cells=cells,
                col_widths_dxa=col_widths_dxa,
                font_name=header_font if is_header else body_font,
                font_sz_half_points=font_sz_half_points,
                bold=is_header,
            )
        )

    return tbl_el


def _document_has_track_changes(doc: "DocxDocument") -> bool:
    """Return True iff ``<w:trackRevisions/>`` is present in settings.xml.

    Mirrors :func:`scitex_msword.track_changes.is_track_changes_enabled`
    without importing it (kept self-contained so callers can use tables
    without track-changes wiring).
    """
    settings = doc.settings.element
    return settings.find(qn("w:trackRevisions")) is not None


def _next_row_revision_id(doc: "DocxDocument") -> int:
    """Pick a w:id one larger than the max existing ``w:ins`` / ``w:del`` id in body."""
    body = doc.element.body
    ins_tag = _w("ins")
    del_tag = _w("del")
    id_attr = qn("w:id")
    max_id = 0
    for el in body.iter():
        if el.tag in (ins_tag, del_tag):
            try:
                eid = int(el.get(id_attr) or "0")
            except (TypeError, ValueError):
                continue
            if eid > max_id:
                max_id = eid
    return max_id + 1


def _now_iso() -> str:
    from datetime import datetime, timezone

    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def insert_table_after_paragraph(
    doc: "DocxDocument",
    paragraph_index: int,
    rows: Sequence[Sequence[str]],
    col_widths_dxa: Sequence[int] = (3000, 6000),
    header_row: bool = True,
    body_font: str = "MS 明朝",
    header_font: str = "MS ゴシック",
    font_size_pt: float = 10.5,
    *,
    track_changes: Optional[bool] = None,
    track_changes_author: str = "agent",
    track_changes_date: Optional[str] = None,
) -> "etree._Element":
    """Insert a Word table immediately after ``doc.paragraphs[paragraph_index]``.

    The table is built directly as a ``<w:tbl>`` element via lxml and
    appended via ``addnext()`` — no python-docx high-level Table API is
    used, so the same code path works on stripped-down ``Document``
    objects where the document part is largely raw OOXML (common in
    Track-Changes-preserving dogfood pipelines).

    Parameters
    ----------
    doc : docx.Document
        The Document to mutate in place.
    paragraph_index : int
        0-based index into ``doc.paragraphs``. The new table is inserted
        as the immediate next sibling of that paragraph in the body.
    rows : sequence of sequence of str
        Row contents. All rows must have the same column count, equal
        to ``len(col_widths_dxa)``. The first row is rendered as a
        header when ``header_row=True``.
    col_widths_dxa : sequence of int, default ``(3000, 6000)``
        Column widths in dxa (twentieths of a point). Length sets the
        column count. Default reproduces the 2-col 1:2 layout used by
        proj-grant's ``build_v43.py``.
    header_row : bool, default ``True``
        When ``True``, row 0 is styled with ``header_font`` + bold; the
        body rows use ``body_font`` without bold.
    body_font : str, default ``"MS 明朝"``
        Font family applied to body-row runs (``rFonts`` ascii / eastAsia
        / hAnsi / cs).
    header_font : str, default ``"MS ゴシック"``
        Font family applied to the header-row runs.
    font_size_pt : float, default ``10.5``
        Run font size in points. Word stores this as half-points in
        ``<w:sz w:val=…>`` (10.5pt → "21").
    track_changes : bool, optional
        Control row-level ``<w:trPr><w:ins/></w:trPr>`` emission:

        - ``None`` (default): auto-detect — wrap iff the document has
          ``<w:trackRevisions/>`` in settings.xml at call time.
        - ``True``: always wrap, regardless of document state.
        - ``False``: never wrap (plain insertion).
    track_changes_author : str, default ``"agent"``
        Author string written to each row's ``w:ins/@w:author``.
    track_changes_date : str, optional
        ISO-8601 timestamp written to each row's ``w:ins/@w:date``.
        Defaults to ``now(UTC)``.

    Returns
    -------
    lxml.etree._Element
        The freshly inserted ``<w:tbl>`` element. Callers that want to
        chain further insertions (e.g. add a paragraph after the table)
        can ``addnext()`` to this element.

    Raises
    ------
    IndexError
        ``paragraph_index`` is out of range for ``doc.paragraphs``.
    ValueError
        ``rows`` is empty, has zero-width rows, or has rows whose column
        count does not match ``len(col_widths_dxa)`` (or, when
        ``col_widths_dxa`` is unset, an inconsistent column count across
        rows).
    ImportError
        ``python-docx`` is not installed.

    Examples
    --------
    .. code-block:: python

        from docx import Document
        from scitex_msword import insert_table_after_paragraph

        doc = Document("draft.docx")
        insert_table_after_paragraph(
            doc,
            paragraph_index=17,
            rows=[
                ("役割", "モジュール名"),
                ("論文執筆", "scitex-writer, scitex-msword"),
            ],
            col_widths_dxa=(3000, 6000),
            header_row=True,
        )
        doc.save("draft_with_table.docx")
    """
    _ensure_docx_available()

    n_paragraphs = len(doc.paragraphs)
    if not -n_paragraphs <= paragraph_index < n_paragraphs:
        raise IndexError(
            f"paragraph_index={paragraph_index} out of range for document "
            f"with {n_paragraphs} paragraphs"
        )

    normalised_rows = _normalise_rows(rows)
    n_cols = len(normalised_rows[0])
    col_widths = list(col_widths_dxa)
    if len(col_widths) != n_cols:
        raise ValueError(
            f"col_widths_dxa has {len(col_widths)} entries but rows have "
            f"{n_cols} columns"
        )

    half_points = _font_size_to_half_points(font_size_pt)

    tbl_el = _build_table(
        rows=normalised_rows,
        col_widths_dxa=col_widths,
        header_row=header_row,
        body_font=body_font,
        header_font=header_font,
        font_sz_half_points=half_points,
    )

    if track_changes is None:
        should_wrap = _document_has_track_changes(doc)
    else:
        should_wrap = bool(track_changes)

    if should_wrap:
        author = track_changes_author
        date = track_changes_date or _now_iso()
        next_id = _next_row_revision_id(doc)
        tr_tag = _w("tr")
        for offset, tr_el in enumerate(tbl_el.findall(tr_tag)):
            _mark_row_as_tracked_insertion(
                tr_el,
                w_id=next_id + offset,
                author=author,
                date=date,
            )

    target_p = doc.paragraphs[paragraph_index]
    target_p._p.addnext(tbl_el)
    return tbl_el


__all__ = ["insert_table_after_paragraph"]
