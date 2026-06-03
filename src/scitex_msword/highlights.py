#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-02 00:00:00
# File: src/scitex_msword/highlights.py
#
# Part of scitex-msword (AGPL-3.0-only). See LICENSE at the repo root.

"""
Visual-mark / highlight utilities for python-docx Documents.

This module is used by the BOOST v16 dogfooding workflow to visualize:

- *additions* by the operator (turquoise highlight, by convention)
- *modifications* to operator-supplied content (magenta highlight)

plus a generic :func:`extract_highlights` to inspect what colors are
present in a document and which runs carry them.

Color names mirror Word's ``WD_COLOR_INDEX`` enum (case-insensitive):
``turquoise``, ``pink``, ``yellow``, ``green``, ``blue``, ``red``,
``gray_25``, ``gray_50``, ``black``, ``white``, ``dark_red``,
``dark_yellow``, ``dark_blue``, ``teal``, ``violet``, and ``auto``.

For convenience the following BOOST-v16 aliases are also accepted:

- ``magenta`` -> ``pink`` (Word's bright-pink highlight)
- ``purple``  -> ``violet``
- ``cyan``    -> ``turquoise``
"""

from __future__ import annotations

from collections import defaultdict
from typing import Any, Dict, Iterable, List, Optional, Tuple

try:
    import docx  # type: ignore[import-untyped]
    from docx.document import Document as DocxDocument  # type: ignore[import-untyped]
    from docx.enum.text import WD_COLOR_INDEX  # type: ignore[import-untyped]

    DOCX_AVAILABLE = True
    _DOCX_IMPORT_ERROR: Optional[Exception] = None
except ImportError as exc:  # pragma: no cover
    DOCX_AVAILABLE = False
    _DOCX_IMPORT_ERROR = exc
    WD_COLOR_INDEX = None  # type: ignore[assignment]
    DocxDocument = None  # type: ignore[assignment,misc]


# Canonical addition / modification colors (mirrors BOOST v16 convention).
ADDITION_COLOR = "turquoise"
MODIFICATION_COLOR = "magenta"

# Map lowercase / underscore-tolerant color name -> WD_COLOR_INDEX value.
# Lazily populated on first use so the module is import-safe without docx.
_COLOR_NAME_MAP: Dict[str, Any] = {}


def _ensure_docx_available() -> None:
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for scitex_msword.highlights. "
            "Install it via `pip install python-docx`."
        ) from _DOCX_IMPORT_ERROR


def _color_map() -> Dict[str, Any]:
    """Lazy-build the lowercase -> WD_COLOR_INDEX mapping."""
    if _COLOR_NAME_MAP:
        return _COLOR_NAME_MAP
    _ensure_docx_available()
    for member in WD_COLOR_INDEX:  # type: ignore[union-attr]
        name = member.name.lower()
        _COLOR_NAME_MAP[name] = member
        # Allow callers to use either "gray25" or "gray_25" forms.
        if "_" in name:
            _COLOR_NAME_MAP[name.replace("_", "")] = member
    # Friendly aliases used by the BOOST v16 workflow.
    _COLOR_NAME_MAP["magenta"] = WD_COLOR_INDEX.PINK  # type: ignore[union-attr]
    _COLOR_NAME_MAP["purple"] = WD_COLOR_INDEX.VIOLET  # type: ignore[union-attr]
    _COLOR_NAME_MAP["cyan"] = WD_COLOR_INDEX.TURQUOISE  # type: ignore[union-attr]
    return _COLOR_NAME_MAP


def _resolve_color(color):
    """Translate a string color name (or pass-through enum) into WD_COLOR_INDEX."""
    if color is None:
        _ensure_docx_available()
        return WD_COLOR_INDEX.AUTO  # type: ignore[union-attr]
    if isinstance(color, str):
        key = color.strip().lower().replace("-", "_")
        cmap = _color_map()
        if key not in cmap:
            raise ValueError(
                f"Unknown highlight color: {color!r}. "
                f"Known: {sorted(set(cmap))[:12]}..."
            )
        return cmap[key]
    return color


def _iter_target_runs(
    document: "DocxDocument", targets: Iterable[Tuple[int, int]]
):
    """Yield (run_obj, (paragraph_idx, run_idx)) pairs for valid targets."""
    paragraphs = list(document.paragraphs)
    for pi, ri in targets:
        if 0 <= pi < len(paragraphs):
            runs = paragraphs[pi].runs
            if 0 <= ri < len(runs):
                yield runs[ri], (pi, ri)


def _apply_highlight(
    document: "DocxDocument",
    runs: Iterable[Tuple[int, int]],
    color,
) -> "DocxDocument":
    """Apply ``color`` highlight to every (paragraph_idx, run_idx) target."""
    _ensure_docx_available()
    wd_color = _resolve_color(color)
    for run, _ in _iter_target_runs(document, runs):
        run.font.highlight_color = wd_color
    return document


def mark_additions(
    document: "DocxDocument",
    runs: Iterable[Tuple[int, int]],
    color: str = ADDITION_COLOR,
) -> "DocxDocument":
    """
    Highlight the runs that the operator (or agent) *added* to the document.

    Parameters
    ----------
    document : docx.Document
        The Document to mutate in place.
    runs : iterable of (paragraph_idx, run_idx)
        Targets to highlight. Out-of-range indices are skipped silently.
    color : str, default "turquoise"
        Color name. See module docstring for the supported palette.

    Returns
    -------
    docx.Document
        The same Document object, mutated.

    Examples
    --------
    >>> from scitex_msword.highlights import mark_additions
    >>> doc = mark_additions(doc, [(3, 0), (5, 2)])  # default turquoise
    """
    return _apply_highlight(document, runs, color)


def mark_modifications(
    document: "DocxDocument",
    runs: Iterable[Tuple[int, int]],
    color: str = MODIFICATION_COLOR,
) -> "DocxDocument":
    """
    Highlight the runs that the operator (or agent) *modified* in the document.

    Parameters
    ----------
    document : docx.Document
        The Document to mutate in place.
    runs : iterable of (paragraph_idx, run_idx)
        Targets to highlight. Out-of-range indices are skipped silently.
    color : str, default "magenta"
        Color name. See module docstring for the supported palette.

    Returns
    -------
    docx.Document
        The same Document object, mutated.

    Examples
    --------
    >>> from scitex_msword.highlights import mark_modifications
    >>> doc = mark_modifications(doc, [(7, 1)])  # default magenta
    """
    return _apply_highlight(document, runs, color)


def _run_highlight_name(run) -> Optional[str]:
    """Return the lowercase highlight color name for the run, or None."""
    try:
        hl = run.font.highlight_color
    except Exception:
        return None
    if hl is None:
        return None
    # WD_COLOR_INDEX members expose .name; str() includes the int code,
    # which is unhelpful for bucketing.
    name = getattr(hl, "name", str(hl)).split(".")[-1].lower()
    # Strip any trailing " (N)" leftover defensively.
    name = name.split(" ")[0]
    if name in ("auto", "inherited"):
        return None
    return name


def extract_highlights(
    document: "DocxDocument",
    by_color: bool = True,
) -> Dict[str, List[Dict[str, Any]]]:
    """
    Extract highlighted runs from a document, grouped by color.

    Parameters
    ----------
    document : docx.Document
        The document to scan.
    by_color : bool, default True
        If True (default), return ``{color_name: [run_info, ...]}``.
        If False, return a single ``{"all": [run_info, ...]}`` bucket
        with each entry's ``color`` field populated.

    Returns
    -------
    dict[str, list[dict]]
        Mapping from color name to a list of run info dicts of shape::

            {"paragraph": int, "run": int, "text": str, "color": str}

    Examples
    --------
    >>> from scitex_msword.highlights import extract_highlights
    >>> by_color = extract_highlights(doc)
    >>> by_color.get("turquoise", [])
    [{'paragraph': 3, 'run': 0, 'text': '...', 'color': 'turquoise'}]
    """
    _ensure_docx_available()
    buckets: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for pi, para in enumerate(document.paragraphs):
        for ri, run in enumerate(para.runs):
            name = _run_highlight_name(run)
            if not name:
                continue
            entry = {
                "paragraph": pi,
                "run": ri,
                "text": run.text,
                "color": name,
            }
            bucket = name if by_color else "all"
            buckets[bucket].append(entry)
    return dict(buckets)


def clear_highlights(
    document: "DocxDocument",
    colors: Optional[Iterable[str]] = None,
) -> "DocxDocument":
    """
    Remove highlights from all runs (optionally only for the listed colors).

    Parameters
    ----------
    document : docx.Document
        Document to mutate in place.
    colors : iterable of str, optional
        If provided, only runs with one of these highlight colors are
        cleared. If ``None`` (default), every highlighted run is cleared.

    Returns
    -------
    docx.Document
        The same Document object, mutated.
    """
    _ensure_docx_available()
    target = None
    if colors is not None:
        target = {c.strip().lower().replace("-", "_") for c in colors}
    for para in document.paragraphs:
        for run in para.runs:
            name = _run_highlight_name(run)
            if name and (target is None or name in target):
                run.font.highlight_color = WD_COLOR_INDEX.AUTO  # type: ignore[union-attr]
    return document


__all__ = [
    "ADDITION_COLOR",
    "MODIFICATION_COLOR",
    "mark_additions",
    "mark_modifications",
    "extract_highlights",
    "clear_highlights",
]
