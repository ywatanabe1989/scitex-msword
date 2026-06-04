#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-02 00:00:00
# File: src/scitex_msword/comments.py
#
# Part of scitex-msword (AGPL-3.0-only). See LICENSE at the repo root.

"""
Comment extraction and (limited) application for python-docx Documents.

Word stores comments in ``word/comments.xml`` inside the .docx ZIP, while
the comment *anchors* (i.e. the ranges the comment refers to) live in
``word/document.xml`` as ``commentRangeStart`` / ``commentRangeEnd``
sibling elements with an ``w:id`` attribute matching the comment.

This module exposes :func:`extract_comments` for reading them back into
a structured list and :func:`apply_comments_as_edits` which honors a
narrow "REPLACE:" comment grammar to perform automated edits.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from xml.etree import ElementTree as ET

try:
    import docx  # type: ignore[import-untyped]
    from docx.document import Document as DocxDocument  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]

    DOCX_AVAILABLE = True
    _DOCX_IMPORT_ERROR: Optional[Exception] = None
except ImportError as exc:  # pragma: no cover
    DOCX_AVAILABLE = False
    _DOCX_IMPORT_ERROR = exc
    DocxDocument = None  # type: ignore[assignment,misc]
    qn = None  # type: ignore[assignment]


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = {"w": _W_NS}

# Grammar for apply_comments_as_edits — intentionally narrow.
_REPLACE_RE = re.compile(r"^\s*REPLACE\s*:\s*(.+?)\s*$", re.IGNORECASE | re.DOTALL)


def _ensure_docx_available() -> None:
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for scitex_msword.comments. "
            "Install it via `pip install python-docx`."
        ) from _DOCX_IMPORT_ERROR


def _docx_path(doc: Union[str, Path, "DocxDocument"]) -> Optional[Path]:
    """Return the on-disk path backing a Document, or None if in-memory."""
    if isinstance(doc, (str, Path)):
        return Path(doc)
    pkg = getattr(doc, "part", None)
    if pkg is None:
        return None
    pkg_pkg = getattr(pkg, "package", None)
    if pkg_pkg is None:
        return None
    pkg_path = getattr(pkg_pkg, "_path", None)
    if pkg_path:
        return Path(pkg_path)
    return None


def _read_comments_xml(source: Union[str, Path, "DocxDocument"]) -> Optional[bytes]:
    """Return the raw word/comments.xml bytes, or None if absent."""
    if isinstance(source, (str, Path)):
        zip_path = Path(source)
    else:
        # In-memory Document: read straight from the OPC part.
        try:
            for part in source.part.package.iter_parts():
                if part.partname == "/word/comments.xml":
                    return part.blob
        except Exception:
            pass
        path = _docx_path(source)
        if path is None:
            return None
        zip_path = path

    if not zip_path.exists():
        return None
    try:
        with zipfile.ZipFile(zip_path) as zf:
            if "word/comments.xml" not in zf.namelist():
                return None
            return zf.read("word/comments.xml")
    except zipfile.BadZipFile:
        return None


def _parse_comments_xml(blob: bytes) -> Dict[str, Dict[str, Any]]:
    """Parse word/comments.xml into ``{comment_id: {author, text, date}}``."""
    root = ET.fromstring(blob)
    comments: Dict[str, Dict[str, Any]] = {}
    for c in root.findall("w:comment", _NS):
        cid = c.get(qn("w:id")) if qn is not None else c.get(f"{{{_W_NS}}}id")
        if cid is None:
            continue
        author = (
            c.get(qn("w:author")) if qn is not None else c.get(f"{{{_W_NS}}}author")
        ) or ""
        date = (
            c.get(qn("w:date")) if qn is not None else c.get(f"{{{_W_NS}}}date")
        ) or ""
        texts: List[str] = []
        for t in c.iter(f"{{{_W_NS}}}t"):
            if t.text:
                texts.append(t.text)
        comments[str(cid)] = {
            "id": int(cid) if str(cid).isdigit() else cid,
            "author": author,
            "date": date,
            "text": "".join(texts),
        }
    return comments


def _scan_anchors(
    document: "DocxDocument",
) -> Dict[str, Dict[str, Any]]:
    """
    Walk the document body and locate comment anchor ranges.

    Returns ``{comment_id_str: {anchor_text, paragraph_range}}``.
    """
    paragraphs = list(document.paragraphs)
    anchors: Dict[str, Dict[str, Any]] = {}
    # Track which comments are currently open and at which paragraph they started.
    open_comments: Dict[str, int] = {}
    # Buffer for in-progress anchor text per comment id.
    open_text: Dict[str, List[str]] = {}

    id_attr = qn("w:id") if qn is not None else f"{{{_W_NS}}}id"

    for pi, para in enumerate(paragraphs):
        # We want a stream of (kind, payload) events from the paragraph XML.
        for elem in para._p.iter():
            tag = elem.tag
            if tag == f"{{{_W_NS}}}commentRangeStart":
                cid = elem.get(id_attr)
                if cid is not None:
                    open_comments[cid] = pi
                    open_text.setdefault(cid, [])
            elif tag == f"{{{_W_NS}}}commentRangeEnd":
                cid = elem.get(id_attr)
                if cid is not None and cid in open_comments:
                    start_p = open_comments.pop(cid)
                    text = "".join(open_text.pop(cid, []))
                    anchors[cid] = {
                        "anchor_text": text,
                        "paragraph_range": [start_p, pi],
                    }
            elif tag == f"{{{_W_NS}}}t":
                # Any text run that runs while a comment is open contributes
                # to that comment's anchor text.
                if elem.text and open_comments:
                    for cid in open_comments:
                        open_text[cid].append(elem.text)
    return anchors


def extract_comments(
    document: Union[str, Path, "DocxDocument"],
) -> List[Dict[str, Any]]:
    """
    Extract Word comments from a .docx file or open Document.

    Parameters
    ----------
    document : str | Path | docx.Document
        Path to the .docx or an already-open Document.

    Returns
    -------
    list[dict]
        One entry per comment::

            {"id": int | str,
             "author": str,
             "date": str,           # ISO timestamp string, may be empty
             "text": str,           # comment body
             "anchor_text": str,    # text the comment is anchored to
             "paragraph_range": [start, end]}

        ``anchor_text`` and ``paragraph_range`` default to ``""`` and
        ``[None, None]`` when no in-document anchor can be located.

    Examples
    --------
    >>> from scitex_msword.comments import extract_comments
    >>> comments = extract_comments("boost-v16.docx")
    >>> [c["text"] for c in comments]
    ['Please rephrase this', 'REPLACE: Use the new wording']
    """
    _ensure_docx_available()
    blob = _read_comments_xml(document)
    if not blob:
        return []
    parsed = _parse_comments_xml(blob)

    # Anchor scanning needs an open Document.
    if isinstance(document, (str, Path)):
        doc_obj = docx.Document(str(document))
    else:
        doc_obj = document
    anchors = _scan_anchors(doc_obj)

    out: List[Dict[str, Any]] = []
    for cid, meta in parsed.items():
        anchor = anchors.get(cid, {})
        out.append(
            {
                "id": meta["id"],
                "author": meta["author"],
                "date": meta["date"],
                "text": meta["text"],
                "anchor_text": anchor.get("anchor_text", ""),
                "paragraph_range": anchor.get("paragraph_range", [None, None]),
            }
        )
    # Sort by numeric id when possible for stable output.
    out.sort(key=lambda d: (isinstance(d["id"], str), d["id"]))
    return out


def _replace_in_paragraph(paragraph, anchor: str, replacement: str) -> bool:
    """Replace ``anchor`` text in ``paragraph`` with ``replacement``.

    Returns True if a replacement occurred.
    """
    text = paragraph.text
    if anchor not in text:
        return False
    new_text = text.replace(anchor, replacement, 1)
    # Rebuild paragraph runs with the new text, keeping first run's formatting.
    first_run_fmt: Dict[str, Any] = {}
    if paragraph.runs:
        r = paragraph.runs[0]
        first_run_fmt = {
            "bold": r.bold,
            "italic": r.italic,
            "underline": r.underline,
        }
    # Clear existing runs.
    for r in list(paragraph._p.findall(qn("w:r"))):
        paragraph._p.remove(r)
    new_run = paragraph.add_run(new_text)
    if first_run_fmt.get("bold") is not None:
        new_run.bold = first_run_fmt["bold"]
    if first_run_fmt.get("italic") is not None:
        new_run.italic = first_run_fmt["italic"]
    if first_run_fmt.get("underline") is not None:
        new_run.underline = first_run_fmt["underline"]
    return True


def apply_comments_as_edits(
    document: "DocxDocument",
    *,
    comments: Optional[List[Dict[str, Any]]] = None,
    grammar: str = "replace",
) -> Dict[str, Any]:
    """
    Apply comments to the document body using a narrow grammar.

    Only the ``REPLACE:`` grammar is currently supported, i.e. a comment
    whose body matches ``r"^\\s*REPLACE\\s*:\\s*(.+?)\\s*$"`` is interpreted
    as "replace this comment's anchor text with the trailing payload".
    Other comments are ignored.

    Parameters
    ----------
    document : docx.Document
        The Document to mutate in place.
    comments : list[dict], optional
        Pre-extracted comments (as returned by :func:`extract_comments`).
        If ``None``, the comments are read from ``document`` directly.
    grammar : str, default "replace"
        Reserved for future expansion. Currently only ``"replace"``
        is recognised.

    Returns
    -------
    dict
        Summary: ``{"applied": int, "skipped": int, "details": [...]}``.

    Examples
    --------
    >>> from scitex_msword.comments import apply_comments_as_edits
    >>> summary = apply_comments_as_edits(doc)
    >>> summary["applied"]
    2
    """
    _ensure_docx_available()
    if grammar != "replace":
        raise ValueError(
            f"Unsupported grammar: {grammar!r}. Only 'replace' is implemented."
        )

    if comments is None:
        comments = extract_comments(document)

    paragraphs = list(document.paragraphs)
    applied = 0
    skipped = 0
    details: List[Dict[str, Any]] = []

    for c in comments:
        m = _REPLACE_RE.match(c.get("text", ""))
        if not m:
            skipped += 1
            details.append({"id": c.get("id"), "status": "skipped", "reason": "no-grammar-match"})
            continue
        replacement = m.group(1)
        anchor = c.get("anchor_text") or ""
        if not anchor:
            skipped += 1
            details.append({"id": c.get("id"), "status": "skipped", "reason": "no-anchor"})
            continue
        start, end = c.get("paragraph_range", [None, None])
        if start is None or end is None:
            skipped += 1
            details.append({"id": c.get("id"), "status": "skipped", "reason": "no-range"})
            continue
        did_apply = False
        for pi in range(max(0, start), min(end + 1, len(paragraphs))):
            if _replace_in_paragraph(paragraphs[pi], anchor, replacement):
                did_apply = True
                break
        if did_apply:
            applied += 1
            details.append({"id": c.get("id"), "status": "applied"})
        else:
            skipped += 1
            details.append(
                {"id": c.get("id"), "status": "skipped", "reason": "anchor-not-found"}
            )

    return {"applied": applied, "skipped": skipped, "details": details}


__all__ = ["extract_comments", "apply_comments_as_edits"]
