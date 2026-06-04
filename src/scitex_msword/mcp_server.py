#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Timestamp: 2026-06-02 00:00:00
# File: src/scitex_msword/mcp_server.py
#
# Part of scitex-msword (AGPL-3.0-only). See LICENSE at the repo root.

"""
MCP server scaffold exposing scitex-msword as a tool surface.

This is a *scaffold*: a focused tool surface is wired up, covering

- legacy BOOST v16 dogfooding tools (diff_docx, mark_additions,
  mark_modifications, preserve_bold_tokens, extract_highlights,
  extract_comments, list_profiles), and
- the full Track-Changes (revision) API surface from PR #8
  (enable_track_changes, is_track_changes_enabled,
  wrap_as_tracked_insertion, wrap_as_tracked_deletion,
  extract_tracked_changes, accept_all_tracked_changes,
  reject_all_tracked_changes).

Further tools can be added by following the ``_tool`` helper pattern.

Install requirements (optional)::

    pip install scitex-msword[mcp]

Run the server (stdio)::

    python -m scitex_msword.mcp_server

The server is implemented against the official Anthropic ``mcp`` Python
SDK using its ``FastMCP`` convenience wrapper. If the ``mcp`` package
is not installed, importing this module still succeeds — only the
``serve()`` entrypoint raises.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

try:
    from mcp.server.fastmcp import FastMCP  # type: ignore[import-not-found]

    MCP_AVAILABLE = True
    _MCP_IMPORT_ERROR: Optional[Exception] = None
except ImportError as exc:  # pragma: no cover — optional dep
    MCP_AVAILABLE = False
    _MCP_IMPORT_ERROR = exc
    FastMCP = None  # type: ignore[assignment,misc]


SERVER_NAME = "scitex-msword"


def _ensure_mcp_available() -> None:
    if not MCP_AVAILABLE:
        raise ImportError(
            "The MCP scaffold requires the 'mcp' package. "
            "Install with: pip install scitex-msword[mcp]"
        ) from _MCP_IMPORT_ERROR


def build_server():  # pragma: no cover — exercised by `serve()` integration
    """
    Construct and return a FastMCP server with scitex-msword tools wired.

    Each tool is a thin wrapper around an existing public API; the
    wrappers translate path-like arguments and serialize results to
    JSON-friendly Python primitives.
    """
    _ensure_mcp_available()

    server = FastMCP(SERVER_NAME)

    # -- Tool: diff_docx ---------------------------------------------------
    @server.tool()
    def diff_docx_tool(a: str, b: str) -> List[Dict[str, Any]]:
        """Diff two .docx files by paragraph (returns list of ops)."""
        from .diff import diff_docx as _diff_docx

        return _diff_docx(Path(a), Path(b))

    # -- Tool: extract_highlights ----------------------------------------
    @server.tool()
    def extract_highlights_tool(path: str) -> Dict[str, List[Dict[str, Any]]]:
        """Extract highlighted runs from a .docx, bucketed by color."""
        import docx as _docx

        from .highlights import extract_highlights as _extract_highlights

        doc = _docx.Document(path)
        return _extract_highlights(doc)

    # -- Tool: extract_comments ------------------------------------------
    @server.tool()
    def extract_comments_tool(path: str) -> List[Dict[str, Any]]:
        """Extract Word comments + their anchors from a .docx."""
        from .comments import extract_comments as _extract_comments

        return _extract_comments(Path(path))

    # -- Tool: mark_additions -------------------------------------------
    @server.tool()
    def mark_additions_tool(
        path: str,
        out: str,
        runs: List[List[int]],
        color: str = "turquoise",
    ) -> str:
        """
        Open ``path``, highlight ``runs`` (list of [paragraph_idx, run_idx]
        pairs), and save to ``out``. Returns the output path string.
        """
        import docx as _docx

        from .highlights import mark_additions as _mark_additions

        doc = _docx.Document(path)
        _mark_additions(doc, [(int(p), int(r)) for p, r in runs], color=color)
        doc.save(out)
        return out

    # -- Tool: mark_modifications ----------------------------------------
    @server.tool()
    def mark_modifications_tool(
        path: str,
        out: str,
        runs: List[List[int]],
        color: str = "magenta",
    ) -> str:
        """Mirror of mark_additions for modifications (default magenta)."""
        import docx as _docx

        from .highlights import mark_modifications as _mark_modifications

        doc = _docx.Document(path)
        _mark_modifications(doc, [(int(p), int(r)) for p, r in runs], color=color)
        doc.save(out)
        return out

    # -- Tool: preserve_bold_tokens --------------------------------------
    @server.tool()
    def preserve_bold_tokens_tool(
        path: str,
        out: str,
        tokens: Sequence[str],
        font_name: str = "MS Gothic",
        case_sensitive: bool = True,
    ) -> str:
        """Bold-emphasize ``tokens`` in ``path``, save to ``out``."""
        import docx as _docx

        from .bold import preserve_bold_tokens as _preserve_bold_tokens

        doc = _docx.Document(path)
        _preserve_bold_tokens(
            doc, list(tokens), font_name=font_name, case_sensitive=case_sensitive
        )
        doc.save(out)
        return out

    # -- Tool: list_profiles ---------------------------------------------
    @server.tool()
    def list_profiles_tool() -> List[str]:
        """List the known scitex-msword profiles."""
        from .profiles import list_profiles as _list_profiles

        return _list_profiles()

    # -- Tool: enable_track_changes --------------------------------------
    @server.tool()
    def enable_track_changes_tool(
        path: str,
        out: str,
        enabled: bool = True,
    ) -> str:
        """Toggle Word's Track Changes switch in ``path``, save to ``out``.

        Idempotent — repeated calls with ``enabled=True`` leave a single
        ``<w:trackChanges/>`` in word/settings.xml. Returns ``out``.
        """
        import docx as _docx

        from .track_changes import enable_track_changes as _enable

        doc = _docx.Document(path)
        _enable(doc, enabled=enabled)
        doc.save(out)
        return out

    # -- Tool: is_track_changes_enabled ----------------------------------
    @server.tool()
    def is_track_changes_enabled_tool(path: str) -> bool:
        """Return True iff Track Changes is on in ``path`` (read-only)."""
        import docx as _docx

        from .track_changes import is_track_changes_enabled as _is_enabled

        doc = _docx.Document(path)
        return _is_enabled(doc)

    # -- Tool: wrap_as_tracked_insertion ---------------------------------
    @server.tool()
    def wrap_as_tracked_insertion_tool(
        path: str,
        out: str,
        paragraph_idx: int,
        run_indices: List[int],
        author: str = "agent",
        date: Optional[str] = None,
    ) -> str:
        """Wrap selected runs of paragraph[``paragraph_idx``] in ``<w:ins>``.

        ``run_indices`` are 0-based positions within the paragraph's runs.
        Word will surface the wrapped content as an accept/reject-able
        insertion authored by ``author``. Returns ``out``.
        """
        import docx as _docx

        from .track_changes import wrap_as_tracked_insertion as _wrap_ins

        doc = _docx.Document(path)
        para = doc.paragraphs[paragraph_idx]
        _wrap_ins(para, list(run_indices), author=author, date=date)
        doc.save(out)
        return out

    # -- Tool: wrap_as_tracked_deletion ----------------------------------
    @server.tool()
    def wrap_as_tracked_deletion_tool(
        path: str,
        out: str,
        paragraph_idx: int,
        run_indices: List[int],
        author: str = "agent",
        date: Optional[str] = None,
    ) -> str:
        """Wrap selected runs of paragraph[``paragraph_idx``] in ``<w:del>``.

        ``run_indices`` are 0-based positions. The wrapped runs' ``<w:t>``
        children become ``<w:delText>`` so Word renders strike-through.
        Returns ``out``.
        """
        import docx as _docx

        from .track_changes import wrap_as_tracked_deletion as _wrap_del

        doc = _docx.Document(path)
        para = doc.paragraphs[paragraph_idx]
        _wrap_del(para, list(run_indices), author=author, date=date)
        doc.save(out)
        return out

    # -- Tool: extract_tracked_changes -----------------------------------
    @server.tool()
    def extract_tracked_changes_tool(path: str) -> List[Dict[str, Any]]:
        """Return every ``<w:ins>``/``<w:del>`` revision as a list of dicts.

        Each entry carries ``{type, paragraph_idx, author, date, id, text}``.
        Read-only — does not modify ``path``.
        """
        import docx as _docx

        from .track_changes import extract_tracked_changes as _extract

        doc = _docx.Document(path)
        return _extract(doc)

    # -- Tool: accept_all_tracked_changes --------------------------------
    @server.tool()
    def accept_all_tracked_changes_tool(path: str, out: str) -> str:
        """Accept all tracked changes in ``path`` (Word "Accept All"), save to ``out``.

        ``<w:ins>`` wrappers unwrapped (content retained); ``<w:del>``
        wrappers removed (content discarded). Returns ``out``.
        """
        import docx as _docx

        from .track_changes import accept_all_tracked_changes as _accept

        doc = _docx.Document(path)
        _accept(doc)
        doc.save(out)
        return out

    # -- Tool: reject_all_tracked_changes --------------------------------
    @server.tool()
    def reject_all_tracked_changes_tool(path: str, out: str) -> str:
        """Reject all tracked changes in ``path`` (Word "Reject All"), save to ``out``.

        ``<w:ins>`` wrappers removed (content discarded); ``<w:del>``
        wrappers unwrapped (content retained, ``<w:delText>`` retagged to
        ``<w:t>``). Returns ``out``.
        """
        import docx as _docx

        from .track_changes import reject_all_tracked_changes as _reject

        doc = _docx.Document(path)
        _reject(doc)
        doc.save(out)
        return out

    return server


def serve() -> None:  # pragma: no cover — IO loop
    """Run the MCP server over stdio (blocking)."""
    _ensure_mcp_available()
    server = build_server()
    server.run()


if __name__ == "__main__":  # pragma: no cover
    serve()


__all__ = ["SERVER_NAME", "build_server", "serve", "MCP_AVAILABLE"]
